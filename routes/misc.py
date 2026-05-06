"""ProspUp — Blueprint Misc.

Routes éparpillées non encore regroupées par domaine :
- Company détail (full / update / events).
- Audit log et activity log.
- Soft-deleted (restore / purge).
- Opportunities (save / delete).
- IA enrichment log + quickadd parsers.
- Métiers (integrate-tags / integrations-cache).
- System (health / version / logs / verify / check-deployment).
- Data export (xlsx / day) + import (save) + rapport-hebdo.
"""
from __future__ import annotations

import datetime
import hashlib
import json
import logging
import os
import re
import sqlite3
import subprocess
import threading
import time
import urllib.error
import urllib.request
from io import BytesIO
from pathlib import Path
from typing import Any, Dict, List, Tuple

from flask import Blueprint, Response, jsonify, request, send_file
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment
from werkzeug.utils import secure_filename

from app import _audit_log, _create_auto_task, log_activity, logger
from config import APP_DIR, APP_VERSION, DATA_DIR
from utils.ai_helpers import _call_ai, _call_ai_web, _load_ai_config, _stream_ai_sse, _stream_ai_web_sse
from utils.auth import _company_owned, _prospect_owned, _uid, login_required, role_required
from utils.common import _now_iso, _today_iso
from utils.db import _conn

misc_bp = Blueprint("misc", __name__)


@misc_bp.get("/api/company/full")
def api_company_full():
    cid = request.args.get("id")
    if not cid:
        return jsonify({"ok": False, "error": "id is required"}), 400
    cid_i = int(cid)
    uid = _uid()
    if not uid:
        return jsonify(ok=False, error="Non authentifié"), 401
    with _conn() as conn:
        company = conn.execute("SELECT * FROM companies WHERE id=? AND owner_id=?;", (cid_i, uid)).fetchone()
        if not company:
            return jsonify({"ok": False, "error": "company not found"}), 404
        prospects = [
            dict(r)
            for r in conn.execute(
                "SELECT * FROM prospects WHERE company_id=? AND owner_id=? ORDER BY id DESC;",
                (cid_i, uid),
            ).fetchall()
        ]
        opps = [
            dict(r)
            for r in conn.execute(
                "SELECT * FROM opportunities WHERE company_id=? ORDER BY COALESCE(updatedAt, createdAt) DESC, id DESC;",
                (cid_i,),
            ).fetchall()
        ]
        # timeline = company_events + push logs of prospects in this company + prospect_events
        events = []
        try:
            rows = conn.execute(
                "SELECT date, type, title, content, meta, createdAt FROM company_events WHERE company_id=? ORDER BY date DESC, id DESC LIMIT 120;",
                (cid_i,),
            ).fetchall()
            for r in rows:
                d = dict(r)
                d["source"] = "company"
                events.append(d)
        except sqlite3.OperationalError as e:
            logger.warning("company_events query failed: %s", e)

        # push logs for prospects
        rows = conn.execute(
            '''
            SELECT l.sentAt AS date, 'push' AS type, 
                   ('Push (' || COALESCE(l.channel,'email') || ')') AS title,
                   COALESCE(l.subject,'') AS content,
                   json_object('to', l.to_email, 'template', l.template_name, 'prospect_id', p.id, 'prospect_name', p.name) AS meta,
                   l.createdAt AS createdAt
            FROM push_logs l
            JOIN prospects p ON p.id = l.prospect_id
            WHERE p.company_id=? AND p.owner_id=?
            ORDER BY l.id DESC
            LIMIT 120;
            ''',
            (cid_i, uid),
        ).fetchall()
        for r in rows:
            events.append(dict(r) | {"source":"push"})

        # prospect events for those prospects
        try:
            rows = conn.execute(
                '''
                SELECT e.date AS date, e.type AS type, e.title AS title, e.content AS content, e.meta AS meta, e.createdAt AS createdAt
                FROM prospect_events e
                JOIN prospects p ON p.id = e.prospect_id
                WHERE p.company_id=? AND p.owner_id=?
                ORDER BY e.date DESC, e.id DESC
                LIMIT 120;
                ''',
                (cid_i, uid),
            ).fetchall()
            for r in rows:
                events.append(dict(r) | {"source":"prospect"})
        except sqlite3.OperationalError as e:
            logger.warning("prospect_events query failed: %s", e)

    # Parse metas and sort
    out_events=[]
    for e in events:
        d=dict(e)
        try:
            d["meta"] = json.loads(d.get("meta") or "null")
        except Exception:
            d["meta"] = d.get("meta")
        out_events.append(d)
    out_events.sort(key=lambda x: str(x.get("date") or x.get("createdAt") or ""), reverse=True)
    return jsonify({"ok": True, "company": dict(company), "prospects": prospects, "opportunities": opps, "timeline": out_events[:200]})


@misc_bp.post("/api/company/update")
def api_company_update():
    payload, err = validate_payload({'id': (str, int)})
    if err:
        return err
    cid = payload.get("id")
    if not cid:
        return jsonify({"ok": False, "error": "id is required"}), 400
    cid_i = int(cid)
    allowed = ["groupe","site","phone","notes","tags","website","linkedin","industry","size","address","city","country","stack","pain_points","budget","urgency"]
    fields = {k: payload.get(k) for k in allowed if k in payload}
    # tags can be list
    if "tags" in fields:
        v = fields["tags"]
        if isinstance(v, list):
            fields["tags"] = json.dumps([str(x).strip() for x in v if str(x).strip()], ensure_ascii=False)
        elif v is None:
            fields["tags"] = "[]"
        else:
            s=str(v).strip()
            if s.startswith("["):
                fields["tags"] = s
            else:
                parts=[t.strip() for t in s.split(",") if t.strip()]
                fields["tags"] = json.dumps(parts, ensure_ascii=False)
    now = _now_iso()
    if not fields:
        return jsonify({"ok": False, "error": "no fields"}), 400
    uid = _uid()
    if not uid:
        return jsonify(ok=False, error="Non authentifié"), 401
    # v23.4: Defensive check — only whitelisted column names can appear in SQL
    _COMPANY_ALLOWED_COLS = frozenset(allowed)
    assert all(k in _COMPANY_ALLOWED_COLS for k in fields), "Invalid column name"
    sets = ", ".join([f"{k}=?" for k in fields.keys()])
    vals = list(fields.values())
    with _conn() as conn:
        conn.execute(f"UPDATE companies SET {sets} WHERE id=? AND owner_id=?;", (*vals, cid_i, uid))
        row = conn.execute("SELECT * FROM companies WHERE id=? AND owner_id=?;", (cid_i, uid)).fetchone()
    
    # Synchroniser si l'entreprise est partagée
    _sync_shared_company_if_needed(cid_i, uid)
    
    _audit_log("update", "company", cid_i, new_value=json.dumps(fields, ensure_ascii=False))
    log_activity('update', 'entreprise', cid_i, row["groupe"] if row else None)
    return jsonify({"ok": True, "company": dict(row) if row else None})


def _sync_shared_company_if_needed(company_id: int, user_id: int) -> None:
    """Synchronise une entreprise partagée si elle est partagée avec d'autres utilisateurs."""
    with _auth_conn() as aconn:
        # Trouver tous les partages pour cette entreprise
        shares = aconn.execute(
            "SELECT from_user_id, to_user_id FROM shared_companies WHERE company_id = ?;",
            (company_id,)
        ).fetchall()
        
        for share in shares:
            from_user_id = share["from_user_id"]
            to_user_id = share["to_user_id"]
            
            # Si l'utilisateur actuel est celui qui a partagé, synchroniser vers le collaborateur
            if user_id == from_user_id:
                _sync_shared_company_to_collaborator(company_id, from_user_id, to_user_id)
            # Si l'utilisateur actuel est le collaborateur, synchroniser vers l'utilisateur source
            elif user_id == to_user_id:
                _sync_shared_company_to_collaborator(company_id, from_user_id, to_user_id)


@misc_bp.get("/api/audit-log")
def api_audit_log():
    """v23.5: Retrieve audit trail. Admin only."""
    user = _get_current_user()
    if not user or user.get("role") != "admin":
        return jsonify(ok=False, error="Admin requis"), 403
    try:
        page = max(1, int(request.args.get("page") or 1))
        limit = min(200, max(1, int(request.args.get("limit") or 50)))
    except (TypeError, ValueError):
        page, limit = 1, 50
    offset = (page - 1) * limit
    entity = request.args.get("entity")
    entity_id = request.args.get("entity_id")
    with _conn() as conn:
        if entity and entity_id:
            rows = conn.execute(
                "SELECT * FROM audit_log WHERE entity=? AND entity_id=? ORDER BY id DESC LIMIT ? OFFSET ?;",
                (entity, int(entity_id), limit, offset)
            ).fetchall()
            total = int(conn.execute("SELECT COUNT(*) FROM audit_log WHERE entity=? AND entity_id=?;", (entity, int(entity_id))).fetchone()[0])
        elif entity:
            rows = conn.execute(
                "SELECT * FROM audit_log WHERE entity=? ORDER BY id DESC LIMIT ? OFFSET ?;",
                (entity, limit, offset)
            ).fetchall()
            total = int(conn.execute("SELECT COUNT(*) FROM audit_log WHERE entity=?;", (entity,)).fetchone()[0])
        else:
            rows = conn.execute("SELECT * FROM audit_log ORDER BY id DESC LIMIT ? OFFSET ?;", (limit, offset)).fetchall()
            total = int(conn.execute("SELECT COUNT(*) FROM audit_log;").fetchone()[0])
    from math import ceil
    return jsonify(ok=True, logs=[dict(r) for r in rows], pagination={"page": page, "limit": limit, "total": total, "pages": ceil(total / limit) if limit else 1})


@misc_bp.get("/api/activity")
@login_required
@role_required('admin')
def api_activity_logs():
    """v27.10: Journal d'activité multi-utilisateurs — admin only."""
    try:
        page = max(1, int(request.args.get("page") or 1))
    except (TypeError, ValueError):
        page = 1
    per_page = 50
    user_id_filter = request.args.get("user_id")
    action_filter = (request.args.get("action") or "").strip()

    where_clauses = []
    params = []
    if user_id_filter:
        try:
            where_clauses.append("user_id = ?")
            params.append(int(user_id_filter))
        except (TypeError, ValueError):
            pass
    if action_filter:
        where_clauses.append("action = ?")
        params.append(action_filter)

    where_sql = ("WHERE " + " AND ".join(where_clauses)) if where_clauses else ""

    with _auth_conn() as conn:
        total = int(conn.execute(
            f"SELECT COUNT(*) AS n FROM activity_logs {where_sql};", params
        ).fetchone()["n"])
        offset = (page - 1) * per_page
        rows = conn.execute(
            f"SELECT * FROM activity_logs {where_sql} ORDER BY created_at DESC LIMIT ? OFFSET ?;",
            params + [per_page, offset]
        ).fetchall()
        users = conn.execute(
            "SELECT DISTINCT user_id, username FROM activity_logs ORDER BY username;"
        ).fetchall()
        action_rows = conn.execute(
            "SELECT DISTINCT action FROM activity_logs ORDER BY action;"
        ).fetchall()

    return jsonify(
        ok=True,
        logs=[dict(r) for r in rows],
        total=total,
        page=page,
        pages=max(1, math.ceil(total / per_page)),
        users=[dict(u) for u in users],
        actions=[a["action"] for a in action_rows]
    )


@misc_bp.post("/api/soft-deleted/restore")
def api_soft_deleted_restore():
    """v23.5: Restore a soft-deleted entity."""
    uid = _uid()
    if not uid:
        return jsonify(ok=False, error="Non authentifié"), 401
    payload = request.get_json(force=True, silent=True) or {}
    entity = payload.get("entity")
    entity_id = payload.get("id")
    if entity not in ("prospect", "company", "candidate") or not entity_id:
        return jsonify(ok=False, error="entity and id required"), 400
    table = {"prospect": "prospects", "company": "companies", "candidate": "candidates"}[entity]
    with _conn() as conn:
        conn.execute(f"UPDATE {table} SET deleted_at=NULL WHERE id=? AND owner_id=?;", (int(entity_id), uid))
    _audit_log("restore", entity, int(entity_id))
    return jsonify(ok=True)


@misc_bp.post("/api/soft-deleted/purge")
def api_soft_deleted_purge():
    """v23.5: Permanently delete items soft-deleted more than 30 days ago. Admin only."""
    user = _get_current_user()
    if not user or user.get("role") != "admin":
        return jsonify(ok=False, error="Admin requis"), 403
    cutoff = (datetime.datetime.now() - datetime.timedelta(days=30)).isoformat(timespec="seconds")
    purged = {}
    with _conn() as conn:
        for tbl in ("prospects", "companies", "candidates"):
            cur = conn.execute(f"DELETE FROM {tbl} WHERE deleted_at IS NOT NULL AND deleted_at < ?;", (cutoff,))
            purged[tbl] = cur.rowcount
    _audit_log("purge", "system", new_value=json.dumps(purged))
    return jsonify(ok=True, purged=purged)


@misc_bp.post("/api/company/events/add")
def api_company_events_add():
    uid = _uid()
    if not uid:
        return jsonify(ok=False, error="Non authentifié"), 401
    payload = request.get_json(force=True, silent=False) or {}
    cid = payload.get("company_id")
    if not cid:
        return jsonify({"ok": False, "error": "company_id is required"}), 400
    with _conn() as conn:
        row = conn.execute("SELECT id FROM companies WHERE id=? AND owner_id=?;", (int(cid), uid)).fetchone()
        if not row:
            return jsonify(ok=False, error="Entreprise non trouvée"), 404
    title = (payload.get("title") or "").strip() or "Note"
    content = (payload.get("content") or "").rstrip()
    etype = (payload.get("type") or "note").strip()
    date = (payload.get("date") or _now_iso()).strip()
    meta = payload.get("meta")
    meta_json = json.dumps(meta, ensure_ascii=False) if meta is not None else None
    now = _now_iso()
    with _conn() as conn:
        conn.execute(
            "INSERT INTO company_events (company_id, date, type, title, content, meta, createdAt) VALUES (?, ?, ?, ?, ?, ?, ?);",
            (int(cid), date, etype, title, content, meta_json, now),
        )
    return jsonify({"ok": True})


@misc_bp.post("/api/opportunities/save")
def api_opportunities_save():
    uid = _uid()
    if not uid:
        return jsonify(ok=False, error="Non authentifié"), 401
    payload = request.get_json(force=True, silent=False) or {}
    cid = payload.get("company_id")
    title = (payload.get("title") or "").strip()
    stage = (payload.get("stage") or "").strip()
    if not cid or not title or not stage:
        return jsonify({"ok": False, "error": "company_id, title, stage are required"}), 400
    with _conn() as conn:
        if not conn.execute("SELECT id FROM companies WHERE id=? AND owner_id=?;", (int(cid), uid)).fetchone():
            return jsonify(ok=False, error="Entreprise non trouvée"), 404
    oid = payload.get("id")
    candidate_name = (payload.get("candidate_name") or "").strip() or None
    candidate_link = (payload.get("candidate_link") or "").strip() or None
    notes = (payload.get("notes") or "").rstrip() or None
    amount = payload.get("amount")
    try:
        amount = float(amount) if amount not in (None, "", "null") else None
    except Exception:
        amount = None
    now = _now_iso()
    with _conn() as conn:
        cur = conn.cursor()
        if oid:
            cur.execute(
                '''
                UPDATE opportunities
                SET title=?, stage=?, candidate_name=?, candidate_link=?, amount=?, notes=?, updatedAt=?
                WHERE id=? AND company_id=?;
                ''',
                (title, stage, candidate_name, candidate_link, amount, notes, now, int(oid), int(cid)),
            )
            if cur.rowcount == 0:
                oid = None
        if not oid:
            cur.execute(
                '''
                INSERT INTO opportunities (company_id, title, stage, candidate_name, candidate_link, amount, notes, createdAt, updatedAt)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?);
                ''',
                (int(cid), title, stage, candidate_name, candidate_link, amount, notes, now, now),
            )
            oid = cur.lastrowid
    return jsonify({"ok": True, "id": oid})


@misc_bp.post("/api/opportunities/delete")
def api_opportunities_delete():
    uid = _uid()
    if not uid:
        return jsonify(ok=False, error="Non authentifié"), 401
    payload = request.get_json(force=True, silent=False) or {}
    oid = payload.get("id")
    if not oid:
        return jsonify({"ok": False, "error": "id is required"}), 400
    with _conn() as conn:
        row = conn.execute("SELECT company_id FROM opportunities WHERE id=?;", (int(oid),)).fetchone()
        if row and conn.execute("SELECT id FROM companies WHERE id=? AND owner_id=?;", (row["company_id"], uid)).fetchone():
            conn.execute("DELETE FROM opportunities WHERE id=?;", (int(oid),))
    return jsonify({"ok": True})


# ====== Prospect quick actions (v6) ======
# Routes /api/prospect/mark_done + /api/prospects/bulk-* + /api/prospects/remove-tag-globally + /api/prospects/update-contacts — voir routes/bulk.py

@misc_bp.post("/api/ia-enrichment-log")
def api_ia_enrichment_log():
    """Log an IA enrichment event to the entity's timeline."""
    uid = _uid()
    if not uid:
        return jsonify(ok=False, error="Non authentifié"), 401
    payload = request.get_json(force=True, silent=True) or {}
    etype = payload.get("type", "")  # prospect, candidate, company
    entity_id = payload.get("entity_id")
    fields_updated = payload.get("fields_updated", "")
    field_count = payload.get("field_count", 0)

    if not entity_id:
        return jsonify(ok=False, error="entity_id required"), 400
    try:
        entity_id_i = int(entity_id)
    except (TypeError, ValueError):
        return jsonify(ok=False, error="entity_id invalide"), 400

    now = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    today = datetime.datetime.now().strftime("%Y-%m-%d")
    title = f"Enrichissement IA — {field_count} champ(s)"
    content = f"Champs mis à jour : {fields_updated}"
    meta = json.dumps({"source": "ia_import", "field_count": field_count}, ensure_ascii=False)

    try:
        with _conn() as conn:
            if etype == "prospect":
                if not _prospect_owned(entity_id_i):
                    return jsonify(ok=False, error="Accès refusé"), 403
                conn.execute(
                    "INSERT INTO prospect_events (prospect_id, date, type, title, content, meta, createdAt) VALUES (?, ?, ?, ?, ?, ?, ?);",
                    (entity_id_i, today, "ia_enrichment", title, content, meta, now),
                )
            elif etype == "company":
                if not _company_owned(entity_id_i):
                    return jsonify(ok=False, error="Accès refusé"), 403
                conn.execute(
                    "INSERT INTO company_events (company_id, date, type, title, content, meta, createdAt) VALUES (?, ?, ?, ?, ?, ?, ?);",
                    (entity_id_i, today, "ia_enrichment", title, content, meta, now),
                )
            elif etype == "candidate":
                if not _candidate_owned(entity_id_i):
                    return jsonify(ok=False, error="Accès refusé"), 403
                # candidates n'ont pas encore de timeline dédiée
            else:
                return jsonify(ok=False, error="type invalide"), 400
        return jsonify(ok=True)
    except Exception as e:
        return jsonify(ok=False, error=str(e)), 500


@misc_bp.post("/api/quickadd/parse-document")
def api_quickadd_parse_document():
    """Extrait le texte d'un PDF ou Word, envoie à Ollama pour identifier prospects/entreprises/candidats, renvoie une liste JSON."""
    uid = _uid()
    if not uid:
        return jsonify(ok=False, error="Non authentifié"), 401
    if "file" not in request.files:
        return jsonify(ok=False, error="Fichier requis"), 400
    entity_type = (request.form.get("entity_type") or "prospect").strip().lower()
    if entity_type not in ("prospect", "company", "candidate"):
        entity_type = "prospect"
    f = request.files["file"]
    if not f or not f.filename:
        return jsonify(ok=False, error="Aucun fichier"), 400
    ok_upload, err_upload = _validate_upload(f, "document")
    if not ok_upload:
        return jsonify(ok=False, error=err_upload[0]), err_upload[1]
    ext = os.path.splitext(f.filename)[1].lower()

    raw = f.read()
    text = ""
    try:
        if ext == ".pdf":
            from pypdf import PdfReader
            reader = PdfReader(BytesIO(raw))
            parts = []
            for page in reader.pages:
                parts.append(page.extract_text() or "")
            text = "\n".join(parts)
        elif ext in (".doc", ".docx"):
            from docx import Document
            doc = Document(BytesIO(raw))
            text = "\n".join(p.text for p in doc.paragraphs)
            for table in doc.tables:
                for row in table.rows:
                    text += "\n" + "\t".join(cell.text.strip() for cell in row.cells)
    except Exception as e:
        logger.exception("Parse document failed: %s", e)
        return jsonify(ok=False, error=f"Impossible de lire le document: {e}"), 400

    text = (text or "").strip()
    if not text or len(text) < 20:
        return jsonify(ok=False, error="Aucun texte extrait ou document trop court."), 400

    # Limiter la taille pour Ollama (éviter timeout)
    if len(text) > 25000:
        text = text[:25000] + "\n[... texte tronqué ...]"

    if entity_type == "prospect":
        prompt = """Tu dois extraire une liste de prospects (contacts B2B : nom, fonction, entreprise, téléphone, email, LinkedIn, notes) à partir du texte ci-dessous.
Retourne UNIQUEMENT un tableau JSON valide, sans texte avant ou après. Chaque élément doit avoir : name (ou nom), fonction (ou function), _company_name (ou entreprise, company), telephone (ou phone), email, linkedin, notes.
Exemple : [{"name":"Jean Dupont","fonction":"Directeur R&D","_company_name":"Acme","telephone":"06...","email":"jean@acme.fr","linkedin":"","notes":""}]
Texte :
"""
    elif entity_type == "company":
        prompt = """Tu dois extraire une liste d'entreprises (nom, site/ville, téléphone, secteur, notes) à partir du texte ci-dessous.
Retourne UNIQUEMENT un tableau JSON valide, sans texte avant ou après. Chaque élément : groupe (ou name, nom), site (ou city), phone (ou telephone), industry (ou sector), notes, tags (tableau de chaînes).
Exemple : [{"groupe":"Acme SA","site":"Paris","phone":"","industry":"Tech","notes":"","tags":[]}]
Texte :
"""
    else:
        prompt = """Tu dois extraire une liste de candidats (nom, rôle, localisation, LinkedIn, téléphone, email, compétences, notes) à partir du texte ci-dessous (CV, liste de profils, etc.).
Retourne UNIQUEMENT un tableau JSON valide, sans texte avant ou après. Chaque élément : name (ou nom), role, location (ou localisation), linkedin, phone (ou telephone), email, skills (tableau de chaînes), sector, notes.
Exemple : [{"name":"Marie Martin","role":"Ingénieur","location":"Lyon","linkedin":"","phone":"","email":"","skills":["Python","Java"],"notes":""}]
Texte :
"""
    prompt += text

    try:
        timeout = min(180, OLLAMA_TIMEOUT + 60)
        raw_response = _call_ai(prompt, timeout=timeout)
        match = re.search(r"\[[\s\S]*\]", raw_response)
        if not match:
            return jsonify(ok=False, error="L'IA n'a pas renvoyé de liste valide. Essayez un modèle plus puissant ou importez en Excel/CSV."), 400
        items = json.loads(match.group(0))
        if not isinstance(items, list):
            items = [items]
        return jsonify(ok=True, items=items, entity_type=entity_type)
    except urllib.error.URLError as e:
        logger.warning("AI unreachable (parse-document): %s", e)
        return jsonify(ok=False, error="IA indisponible. Vérifiez la configuration dans Paramètres > Configuration IA."), 503
    except json.JSONDecodeError as e:
        logger.warning("AI invalid JSON (parse-document): %s", e)
        return jsonify(ok=False, error="Réponse IA invalide (modèle peut-être trop léger). Essayez un modèle plus puissant ou importez en Excel/CSV."), 400
    except Exception as e:
        logger.exception("quickadd parse-document failed: %s", e)
        return jsonify(ok=False, error=str(e)), 500


def _sse_message(event: str, data: Any) -> str:
    """Format one SSE message (event + data). data can be dict (will be JSON-encoded) or str."""
    payload = json.dumps(data, ensure_ascii=False) if isinstance(data, dict) else str(data)
    return f"event: {event}\ndata: {payload}\n\n"


@misc_bp.post("/api/quickadd/parse-document-stream")
def api_quickadd_parse_document_stream():
    """Like parse-document but streams SSE: phase (upload, extract, ollama), then token events, then done with items.
    Allows the client to show live progress. File must be in request.files['file'], entity_type in form."""
    uid = _uid()
    if not uid:
        return Response(_sse_message("error", {"message": "Non authentifié"}), status=401, mimetype="text/event-stream")
    if "file" not in request.files:
        return Response(_sse_message("error", {"message": "Fichier requis"}), status=400, mimetype="text/event-stream")
    entity_type = (request.form.get("entity_type") or "prospect").strip().lower()
    if entity_type not in ("prospect", "company", "candidate"):
        entity_type = "prospect"
    f = request.files["file"]
    if not f or not f.filename:
        return Response(_sse_message("error", {"message": "Aucun fichier"}), status=400, mimetype="text/event-stream")
    ok_upload, err_upload = _validate_upload(f, "document")
    if not ok_upload:
        return Response(_sse_message("error", {"message": err_upload[0]}), status=err_upload[1], mimetype="text/event-stream")
    ext = os.path.splitext(f.filename)[1].lower()

    def generate():
        try:
            yield _sse_message("phase", {"step": "extract", "label": "Extraction du document…"})
            raw = f.read()
            text = ""
            if ext == ".pdf":
                from pypdf import PdfReader
                reader = PdfReader(BytesIO(raw))
                text = "\n".join((p.extract_text() or "") for p in reader.pages)
            elif ext in (".doc", ".docx"):
                from docx import Document
                doc = Document(BytesIO(raw))
                text = "\n".join(p.text for p in doc.paragraphs)
                for table in doc.tables:
                    for row in table.rows:
                        text += "\n" + "\t".join(cell.text.strip() for cell in row.cells)
            text = (text or "").strip()
            if not text or len(text) < 20:
                yield _sse_message("error", {"message": "Aucun texte extrait ou document trop court."})
                return
            if len(text) > 25000:
                text = text[:25000] + "\n[... texte tronqué ...]"

            if entity_type == "prospect":
                prompt = """Tu dois extraire une liste de prospects (contacts B2B : nom, fonction, entreprise, téléphone, email, LinkedIn, notes) à partir du texte ci-dessous.
Retourne UNIQUEMENT un tableau JSON valide, sans texte avant ou après. Chaque élément doit avoir : name (ou nom), fonction (ou function), _company_name (ou entreprise, company), telephone (ou phone), email, linkedin, notes.
Exemple : [{"name":"Jean Dupont","fonction":"Directeur R&D","_company_name":"Acme","telephone":"06...","email":"jean@acme.fr","linkedin":"","notes":""}]
Texte :
"""
            elif entity_type == "company":
                prompt = """Tu dois extraire une liste d'entreprises (nom, site/ville, téléphone, secteur, notes) à partir du texte ci-dessous.
Retourne UNIQUEMENT un tableau JSON valide, sans texte avant ou après. Chaque élément : groupe (ou name, nom), site (ou city), phone (ou telephone), industry (ou sector), notes, tags (tableau de chaînes).
Exemple : [{"groupe":"Acme SA","site":"Paris","phone":"","industry":"Tech","notes":"","tags":[]}]
Texte :
"""
            else:
                prompt = """Tu dois extraire une liste de candidats (nom, rôle, localisation, LinkedIn, téléphone, email, compétences, notes) à partir du texte ci-dessous (CV, liste de profils, etc.).
Retourne UNIQUEMENT un tableau JSON valide, sans texte avant ou après. Chaque élément : name (ou nom), role, location (ou localisation), linkedin, phone (ou telephone), email, skills (tableau de chaînes), sector, notes.
Exemple : [{"name":"Marie Martin","role":"Ingénieur","location":"Lyon","linkedin":"","phone":"","email":"","skills":["Python","Java"],"notes":""}]
Texte :
"""
            prompt += text

            config = _load_ai_config()
            provider_label = "Groq" if config.get("provider") == "groq" else "Ollama"
            yield _sse_message("phase", {"step": "ollama", "label": f"Analyse par l'IA ({provider_label})…"})
            timeout = min(180, OLLAMA_TIMEOUT + 60)
            full_response = []
            for sse_line in _stream_ai_sse(prompt, None, timeout):
                if not sse_line.startswith("data: "):
                    continue
                data_str = sse_line.strip().removeprefix("data: ").strip()
                if not data_str:
                    continue
                try:
                    evt = json.loads(data_str)
                except json.JSONDecodeError:
                    continue
                if evt.get("type") == "token":
                    token_text = evt.get("text", "")
                    if token_text:
                        full_response.append(token_text)
                        yield _sse_message("token", {"text": token_text})
                elif evt.get("type") == "error":
                    yield _sse_message("error", {"message": evt.get("message", "Erreur IA")})
                    return
            raw_response = "".join(full_response)
            match = re.search(r"\[[\s\S]*\]", raw_response)
            if not match:
                yield _sse_message("error", {
                    "message": "L'IA n'a pas renvoyé de liste valide. Essayez un modèle plus puissant ou importez en Excel/CSV."
                })
                return
            try:
                items = json.loads(match.group(0))
            except json.JSONDecodeError:
                yield _sse_message("error", {"message": "Réponse IA invalide. Essayez un modèle plus puissant ou importez en Excel/CSV."})
                return
            if not isinstance(items, list):
                items = [items]
            yield _sse_message("done", {"items": items, "entity_type": entity_type})
        except urllib.error.URLError as e:
            logger.warning("AI unreachable (parse-document-stream): %s", e)
            yield _sse_message("error", {
                "message": "IA indisponible. Vérifiez la configuration dans Paramètres > Configuration IA."
            })
        except Exception as e:
            logger.exception("quickadd parse-document-stream failed: %s", e)
            yield _sse_message("error", {"message": str(e)})

    return Response(
        generate(),
        mimetype="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


# /api/ollama/generate, /generate-stream, /api/ai/config, /api/ai/test
# — déplacés dans routes/ai.py

# ═══════════════════════════════════════════════════════════════════
# v25.8: Intégration automatique des tags dans l'arbre des métiers via Ollama
# ═══════════════════════════════════════════════════════════════════
_TAG_INTEGRATION_CACHE_FILE = APP_DIR / "data" / "tag_integrations.json"

def _load_tag_integrations() -> Dict[str, Dict[str, Any]]:
    """Charge le cache des intégrations de tags."""
    if _TAG_INTEGRATION_CACHE_FILE.exists():
        try:
            with open(_TAG_INTEGRATION_CACHE_FILE, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception as e:
            logger.warning("Erreur chargement cache intégrations tags: %s", e)
    return {}

def _save_tag_integrations(cache: Dict[str, Dict[str, Any]]):
    """Sauvegarde le cache des intégrations de tags."""
    try:
        _TAG_INTEGRATION_CACHE_FILE.parent.mkdir(parents=True, exist_ok=True)
        with open(_TAG_INTEGRATION_CACHE_FILE, "w", encoding="utf-8") as f:
            json.dump(cache, f, ensure_ascii=False, indent=2)
    except Exception as e:
        logger.error("Erreur sauvegarde cache intégrations tags: %s", e)

@misc_bp.post("/api/metiers/integrate-tags")
def api_metiers_integrate_tags():
    """Intègre automatiquement des tags manquants dans l'arbre des métiers via Ollama.
    
    Phase 1 amélioré : utilise aussi la similarité sémantique pour trouver les meilleures correspondances.
    
    Reçoit: { "tags": ["tag1", "tag2"], "context": { "company": "...", "fonction": "...", "linkedin": "..." } }
    Retourne: { "ok": true, "integrations": { "tag1": { "category": "...", "specialty": "...", "techCategory": "...", "similarity": 0.85 } } }
    """
    uid = _uid()
    if not uid:
        return jsonify(ok=False, error="Non authentifié"), 401
    
    payload = request.get_json(force=True, silent=True) or {}
    tags = payload.get("tags", [])
    context = payload.get("context", {})
    
    if not tags or not isinstance(tags, list):
        return jsonify(ok=False, error="Liste de tags requise"), 400
    
    cache = _load_tag_integrations()
    results = {}
    
    # Structure des métiers pour le prompt Ollama
    metiers_structure_detailed = """Ingénierie Logicielle:
  - Logiciel applicatif
  - Test / Validation / Qualification logicielle
  - Logiciels embarqués / Systèmes embarqués / IoT
  - Data Science / ML / Deep Learning / Vision
  - DevOps / Infrastructure / Cloud
  - Gestion de projet logiciel / Scrum Master
  - Développement Web / Fullstack

Ingénierie Électronique:
  - Électronique analogique
  - Électronique numérique
  - Électronique de puissance
  - Génie électrique / Électrotechnique
  - Industrialisation
  - FPGA / ASIC / SoC

Ingénierie Système:
  - Mécatronique / Robotique
  - Model Based Design (MBD)
  - Safety / Sûreté de fonctionnement
  - Contrôle commande / Automatique
  - Simulation multiphysique / Modélisation
  - Mécanique
  - Système (ingénierie système)
  - Test / Validation / Essais système

Life Science:
  - Qualification d'équipements (Pharma & DM)
  - Validation de systèmes automatisés (VSA)
  - Validation de systèmes d'informations (VSI)
  - Validation de produits (Dispositifs Médicaux)"""
    
    # Liste des catégories de tech possibles
    tech_categories = [
        "Langages", "Systèmes", "IDE", "Bases de données", "Méthodologies",
        "Outils", "Librairies", "Protocoles", "Microcontrôleurs", "Capteurs",
        "Frameworks", "Matériel", "Outils CAO", "Serveurs", "Secteurs"
    ]
    
    for tag in tags:
        tag_lower = tag.lower().strip()
        
        # Vérifier le cache
        if tag_lower in cache:
            results[tag] = cache[tag_lower]
            continue
        
        # Construire le prompt pour Ollama
        context_str = ""
        if context.get("company"):
            context_str += f"Entreprise: {context['company']}. "
        if context.get("fonction"):
            context_str += f"Poste: {context['fonction']}. "
        if context.get("linkedin"):
            context_str += f"LinkedIn disponible. "
        
        prompt = f"""Tu es un expert en classification de compétences techniques pour l'ingénierie.

Contexte du prospect: {context_str}

Tag à classer: "{tag}"

Arbre des métiers disponible:
{metiers_structure_detailed}

Catégories de technologies possibles: {', '.join(tech_categories)}

Instructions:
1. Analyse le tag "{tag}" dans le contexte donné
2. Identifie la catégorie métier (Ingénierie Logicielle, Ingénierie Électronique, Ingénierie Système, ou Life Science)
3. Identifie la spécialité la plus appropriée dans cette catégorie
4. Identifie la catégorie de technologie la plus appropriée

Réponds UNIQUEMENT avec un JSON valide au format suivant (sans markdown, sans code block):
{{"category": "Nom exact de la catégorie métier", "specialty": "Nom exact de la spécialité", "techCategory": "Catégorie de technologie la plus appropriée", "reasoning": "Explication courte (1 phrase)"}}

Si le tag ne correspond clairement à aucun métier, réponds avec {{"category": null, "reasoning": "..."}}."""
        
        # Phase 1: Essayer d'abord la similarité sémantique avec les tags du référentiel
        # Charger tous les tags du référentiel depuis metiers-data.js (via import ou lecture)
        # Pour l'instant, on utilise Ollama directement mais on pourrait améliorer avec embeddings
        
        try:
            response_text = _call_ai(prompt, timeout=60)
            
            # Extraire le JSON de la réponse (gérer les blocs de code markdown)
            json_block = re.search(r'```(?:json)?\s*(\{.*?\})\s*```', response_text, re.DOTALL)
            if json_block:
                response_text = json_block.group(1)
            else:
                # Chercher directement un objet JSON
                json_match = re.search(r'\{[^{}]*(?:\{[^{}]*\}[^{}]*)*\}', response_text, re.DOTALL)
                if json_match:
                    response_text = json_match.group(0)
            
            try:
                integration = json.loads(response_text)
                if integration.get("category") and integration.get("specialty") and integration.get("category") != "null":
                    # Phase 1: Calculer similarité avec tags référentiel (optionnel, pour info)
                    # On pourrait améliorer en comparant avec les tags existants dans la spécialité trouvée
                    cache[tag_lower] = integration
                    results[tag] = integration
                else:
                    results[tag] = {"category": None, "reason": "Tag non classable selon Ollama"}
            except json.JSONDecodeError:
                results[tag] = {"category": None, "reason": "Réponse Ollama invalide (JSON non parsable)"}
        except urllib.error.URLError:
            results[tag] = {"category": None, "reason": "Ollama indisponible"}
        except Exception as e:
            logger.warning("Erreur intégration tag %s: %s", tag, e)
            results[tag] = {"category": None, "reason": str(e)}
    
    # Sauvegarder le cache
    if results:
        _save_tag_integrations(cache)
    
    return jsonify(ok=True, integrations=results)


@misc_bp.get("/api/metiers/integrations-cache")
def api_metiers_integrations_cache():
    """Retourne le cache des intégrations de tags."""
    uid = _uid()
    if not uid:
        return jsonify(ok=False, error="Non authentifié"), 401
    cache = _load_tag_integrations()
    return jsonify(ok=True, integrations=cache)


# ═══════════════════════════════════════════════════════════════════
# Post-update validation — rollback automatique si non confirmé en 3 min
# ═══════════════════════════════════════════════════════════════════
_VALIDATION_TIMER: threading.Timer | None = None
_VALIDATION_LOCK = threading.Lock()
_VALIDATION_TIMEOUT_SECONDS = 180  # 3 minutes


def _write_pending_validation(previous_commit_full: str) -> None:
    """Écrit .pending_validation avant le restart post-pull."""
    data = {
        "triggered_at": datetime.datetime.now().isoformat(timespec="seconds"),
        "previous_commit": previous_commit_full,
        "timeout_seconds": _VALIDATION_TIMEOUT_SECONDS,
    }
    try:
        (APP_DIR / ".pending_validation").write_text(
            json.dumps(data, ensure_ascii=False), encoding="utf-8"
        )
        logger.info("Pending validation écrit (rollback vers %s si non confirmé dans %ds)",
                    previous_commit_full[:7], _VALIDATION_TIMEOUT_SECONDS)
    except Exception as e:
        logger.warning("Impossible d'écrire .pending_validation: %s", e)


def _auto_rollback_on_timeout() -> None:
    """Déclenché après 3 min sans confirmation — rollback automatique vers le commit précédent."""
    pv_file = APP_DIR / ".pending_validation"
    if not pv_file.exists():
        return  # Déjà confirmé ou annulé

    logger.warning("Validation timeout — rollback automatique déclenché")

    previous_commit = ""
    triggered_at = ""
    try:
        pv_data = json.loads(pv_file.read_text(encoding="utf-8"))
        previous_commit = pv_data.get("previous_commit", "")
        triggered_at = pv_data.get("triggered_at", "")
    except Exception:
        pass

    # Journal d'erreur détaillé
    error_log: dict = {
        "reason": "timeout",
        "message": "Aucune confirmation reçue dans les 3 minutes — rollback automatique",
        "triggered_at": triggered_at,
        "rollback_at": datetime.datetime.now().isoformat(timespec="seconds"),
        "previous_commit": previous_commit[:7] if previous_commit else "unknown",
        "git_log": "",
    }
    try:
        cp_log = subprocess.run(
            ["git", "log", "--oneline", "-10"],
            cwd=str(APP_DIR), capture_output=True, text=True, timeout=5,
        )
        error_log["git_log"] = cp_log.stdout.strip()
    except Exception:
        pass

    try:
        (APP_DIR / ".validation_error_log").write_text(
            json.dumps(error_log, ensure_ascii=False, indent=2), encoding="utf-8"
        )
    except Exception as e:
        logger.error("Impossible d'écrire .validation_error_log: %s", e)

    try:
        pv_file.unlink(missing_ok=True)
    except Exception:
        pass

    # Rollback git vers le commit précédent
    if previous_commit:
        try:
            result = subprocess.run(
                ["git", "reset", "--hard", previous_commit],
                cwd=str(APP_DIR), capture_output=True, text=True, timeout=30,
            )
            logger.info("Rollback auto: git reset --hard %s → returncode=%d",
                        previous_commit[:7], result.returncode)
        except Exception as e:
            logger.error("Rollback auto git échoué: %s", e)

    _schedule_restart(delay=3.0)


def _start_validation_timer() -> None:
    """Lance le timer de validation (3 min avant rollback automatique)."""
    global _VALIDATION_TIMER
    with _VALIDATION_LOCK:
        if _VALIDATION_TIMER is not None:
            _VALIDATION_TIMER.cancel()
        _VALIDATION_TIMER = threading.Timer(_VALIDATION_TIMEOUT_SECONDS, _auto_rollback_on_timeout)
        _VALIDATION_TIMER.daemon = True
        _VALIDATION_TIMER.start()
    logger.info("Timer de validation démarré (%ds)", _VALIDATION_TIMEOUT_SECONDS)


def _cancel_validation_timer() -> None:
    """Annule le timer de validation (appelé quand l'utilisateur confirme)."""
    global _VALIDATION_TIMER
    with _VALIDATION_LOCK:
        if _VALIDATION_TIMER is not None:
            _VALIDATION_TIMER.cancel()
            _VALIDATION_TIMER = None
    logger.info("Timer de validation annulé (confirmation reçue)")


def _schedule_restart(delay: float = 10.0):
    """Restart after responding.

    - If launched via PROSPUP.bat (or _run_serveur.bat), it will restart on exit code 42.
    - If launched directly (python app.py), it spawns a new process then exits.
    - On Windows, the new process is detached to survive terminal closure (Cursor, etc.).

    Le délai permet aux clients (Cloudflare, navigateurs) de recevoir la réponse HTTP
    avant que le serveur ne redémarre, évitant les erreurs 502.
    """
    def _do():
        time.sleep(float(delay))
        launcher = (os.environ.get("PROSPUP_LAUNCHER") or "").strip().upper()
        if launcher == "BAT":
            logger.info("Restart: exit code 42 pour le superviseur")
            os._exit(42)
        try:
            import sys as _sys
            args = [_sys.executable] + _sys.argv
            logger.info("Restart: lancement nouveau processus: %s", " ".join(args))
            
            # Sur Windows, détacher le processus pour qu'il survive à la fermeture du terminal
            # (utile quand lancé depuis Cursor ou un terminal qui peut être fermé)
            creation_flags = 0
            if sys.platform == "win32":
                # CREATE_NEW_PROCESS_GROUP + DETACHED_PROCESS pour indépendance du terminal
                # Note: DETACHED_PROCESS peut ne pas être disponible sur toutes les versions
                try:
                    creation_flags = subprocess.CREATE_NEW_PROCESS_GROUP | subprocess.DETACHED_PROCESS
                except AttributeError:
                    # Fallback si DETACHED_PROCESS n'existe pas (anciennes versions Python)
                    creation_flags = subprocess.CREATE_NEW_PROCESS_GROUP
            
            proc = subprocess.Popen(
                args,
                cwd=str(APP_DIR),
                creationflags=creation_flags if sys.platform == "win32" else 0,
                # Sur Unix, utiliser start_new_session pour détacher du terminal
                start_new_session=(sys.platform != "win32")
            )
            time.sleep(2.0)
            logger.info("Restart: nouveau processus lancé (PID %d), arrêt de l'ancien serveur", proc.pid)
        except Exception as e:
            logger.error("Restart: erreur lors du lancement du nouveau processus: %s", e)
        os._exit(0)

    threading.Thread(target=_do, daemon=True).start()


# /api/deploy/pull, /restart, /health, /validation-status, /confirm-validation
# — déplacés dans routes/deploy.py


@misc_bp.route("/api/system/check-deployment", methods=["GET"])
def api_system_check_deployment():
    """Vérifie si le code de vérification système est déployé."""
    user = _get_current_user()
    if not user or user.get("role") != "admin":
        return jsonify(ok=False, error="Admin requis"), 403
    
    verify_script = APP_DIR / "scripts" / "verify_all.py"
    verify_script_exists = verify_script.exists()
    
    # Vérifier si la section est dans templates/parametres.html
    parametres_file = APP_DIR / "templates" / "parametres.html"
    has_section = False
    if parametres_file.exists():
        try:
            content = parametres_file.read_text(encoding="utf-8")
            has_section = "systemVerifySection" in content and "Vérification système" in content
        except Exception:
            pass
    
    # Vérifier aussi si le fichier existe à la racine (compatibilité)
    if not has_section:
        parametres_file_root = APP_DIR / "parametres.html"
        if parametres_file_root.exists():
            try:
                content = parametres_file_root.read_text(encoding="utf-8")
                has_section = "systemVerifySection" in content and "Vérification système" in content
            except Exception:
                pass
    
    # Vérifier si la fonction JS existe
    page_settings_file = APP_DIR / "static" / "js" / "page-settings.js"
    has_js_function = False
    if page_settings_file.exists():
        try:
            content = page_settings_file.read_text(encoding="utf-8")
            has_js_function = "runSystemVerify" in content
        except Exception:
            pass
    
    # Dernier commit et branche (pour affichage "version en ligne")
    last_commit = "unknown"
    commit_hash = "unknown"
    branch = "main"
    try:
        cp = subprocess.run(
            ["git", "log", "-1", "--oneline", "HEAD"],
            cwd=str(APP_DIR),
            capture_output=True,
            text=True,
            timeout=2,
        )
        if cp.returncode == 0:
            last_commit = (cp.stdout or "").strip()[:50]
        cp2 = subprocess.run(
            ["git", "rev-parse", "HEAD"],
            cwd=str(APP_DIR),
            capture_output=True,
            text=True,
            timeout=2,
        )
        if cp2.returncode == 0:
            commit_hash = (cp2.stdout or "").strip()[:7]
        cp3 = subprocess.run(
            ["git", "rev-parse", "--abbrev-ref", "HEAD"],
            cwd=str(APP_DIR),
            capture_output=True,
            text=True,
            timeout=2,
        )
        if cp3.returncode == 0 and (cp3.stdout or "").strip():
            branch = (cp3.stdout or "").strip()
    except Exception:
        pass
    
    return jsonify(
        ok=True,
        verify_script_exists=verify_script_exists,
        html_section_exists=has_section,
        js_function_exists=has_js_function,
        all_deployed=verify_script_exists and has_section and has_js_function,
        last_commit=last_commit,
        version=APP_VERSION,
        commit_hash=commit_hash,
        branch=branch,
    )


@misc_bp.route("/api/system/logs", methods=["GET"])
def api_system_logs():
    """Retourne les dernières lignes du log serveur. Admin uniquement."""
    user = _get_current_user()
    if not user or user.get("role") != "admin":
        return jsonify(ok=False, error="Admin requis"), 403
    
    log_file = APP_DIR / "logs" / "prospup.log"
    lines = request.args.get("lines", 50, type=int)
    lines = min(max(10, lines), 500)  # Entre 10 et 500 lignes
    
    if not log_file.exists():
        return jsonify(ok=False, error="Fichier de log introuvable"), 404
    
    try:
        # Lire les dernières lignes du fichier
        with open(log_file, "r", encoding="utf-8", errors="ignore") as f:
            all_lines = f.readlines()
            last_lines = all_lines[-lines:] if len(all_lines) > lines else all_lines
        
        return jsonify(
            ok=True,
            lines=last_lines,
            total_lines=len(all_lines),
            file_size=log_file.stat().st_size,
        )
    except Exception as e:
        logger.exception("Failed to read logs")
        return jsonify(ok=False, error=str(e)), 500


@misc_bp.post("/api/system/verify")
def api_system_verify():
    """Exécute le script de vérification système et retourne les résultats détaillés."""
    user = _get_current_user()
    if not user or user.get("role") != "admin":
        return jsonify(ok=False, error="Admin requis"), 403
    
    verify_script = APP_DIR / "scripts" / "verify_all.py"
    if not verify_script.exists():
        return jsonify(ok=False, error="Script de vérification introuvable"), 404
    
    try:
        # Exécuter le script avec capture de la sortie
        proc = subprocess.run(
            [sys.executable, str(verify_script)],
            cwd=str(APP_DIR),
            capture_output=True,
            text=True,
            timeout=60,
        )
        
        # Parser les résultats (le script utilise des exit codes)
        checks = {
            "git": {"ok": True, "message": "OK"},
            "ollama": {"ok": True, "message": "OK"},
            "flask": {"ok": True, "message": "OK"},
            "api_ollama": {"ok": True, "message": "OK"},
            "scripts": {"ok": True, "message": "OK"},
            "env": {"ok": True, "message": "OK"},
        }
        
        # Déterminer quel check a échoué selon l'exit code
        if proc.returncode == 1:
            checks["git"]["ok"] = False
            checks["git"]["message"] = proc.stderr or "Erreur Git (repo, branche ou pull)"
        elif proc.returncode == 2:
            checks["ollama"]["ok"] = False
            checks["ollama"]["message"] = proc.stderr or "Ollama inaccessible ou modèle introuvable"
        elif proc.returncode == 3:
            checks["flask"]["ok"] = False
            checks["flask"]["message"] = proc.stderr or "Flask ne répond pas"
        elif proc.returncode == 4:
            checks["api_ollama"]["ok"] = False
            checks["api_ollama"]["message"] = proc.stderr or "API Ollama via Flask en erreur (possible erreur 405)"
        elif proc.returncode == 5:
            checks["scripts"]["ok"] = False
            checks["scripts"]["message"] = proc.stderr or "Erreur dans les scripts Python"
        elif proc.returncode == 6:
            checks["env"]["ok"] = False
            checks["env"]["message"] = proc.stderr or "Variables d'environnement invalides"
        
        all_ok = proc.returncode == 0
        
        return jsonify(
            ok=all_ok,
            exit_code=proc.returncode,
            checks=checks,
            stdout=proc.stdout,
            stderr=proc.stderr,
        )
    except subprocess.TimeoutExpired:
        return jsonify(ok=False, error="Timeout lors de l'exécution du script"), 504
    except Exception as e:
        logger.exception("System verify failed")
        return jsonify(ok=False, error=str(e)), 500


@misc_bp.route("/api/app-version", methods=["GET"])
def api_app_version():
    """Retourne la version de l'app, le hash du commit et la date du dernier commit pour affichage badge."""
    try:
        # Hash du commit actuel
        cp = subprocess.run(
            ["git", "rev-parse", "HEAD"],
            cwd=str(APP_DIR),
            capture_output=True,
            text=True,
            timeout=2,
        )
        commit_hash = (cp.stdout or "").strip()[:7] if cp.returncode == 0 else "unknown"
        
        # Date du dernier commit
        cp2 = subprocess.run(
            ["git", "log", "-1", "--format=%ci", "HEAD"],
            cwd=str(APP_DIR),
            capture_output=True,
            text=True,
            timeout=2,
        )
        commit_date = (cp2.stdout or "").strip() if cp2.returncode == 0 else ""
        
        # Branche actuelle (ex. main)
        cp3 = subprocess.run(
            ["git", "rev-parse", "--abbrev-ref", "HEAD"],
            cwd=str(APP_DIR),
            capture_output=True,
            text=True,
            timeout=2,
        )
        branch = (cp3.stdout or "").strip() or "main"
        
        # Générer une couleur basée sur le hash (pour changement visuel)
        if commit_hash != "unknown":
            # Utiliser les 6 premiers caractères du hash pour générer une couleur
            hash_int = int(commit_hash[:6], 16) if len(commit_hash) >= 6 else 0
            # Palette de couleurs vives mais lisibles
            colors = [
                "#3b82f6", "#22c55e", "#f59e0b", "#ef4444", "#8b5cf6",
                "#ec4899", "#14b8a6", "#6366f1", "#f97316", "#06b6d4"
            ]
            color_index = hash_int % len(colors)
            badge_color = colors[color_index]
        else:
            badge_color = "#64748b"
        
        return jsonify(ok=True, version=APP_VERSION, commit_hash=commit_hash, commit_date=commit_date, branch=branch, badge_color=badge_color)
    except Exception as e:
        logger.warning("App version fetch error: %s", e)
        return jsonify(ok=True, version=APP_VERSION, commit_hash="unknown", commit_date="", branch="main", badge_color="#64748b")


@misc_bp.get("/api/health")
def api_health():
    """Health check endpoint. Sensitive details only for admins (v23.4)."""
    info: Dict[str, Any] = {"status": "ok", "version": APP_VERSION}
    is_admin = False
    user = _get_current_user()
    if user and user.get('role') == 'admin':
        is_admin = True

    _table_names = ("prospects", "companies", "push_logs", "candidates")
    try:
        with _conn() as con:
            cur = con.cursor()
            for tbl in _table_names:
                try:
                    info[f"{tbl}_count"] = int(cur.execute(
                        f"SELECT COUNT(*) FROM {tbl}"  # noqa: table names are hardcoded above
                    ).fetchone()[0])
                except Exception:
                    info[f"{tbl}_count"] = None
    except Exception as e:
        info["db_error"] = "unavailable"
        if is_admin:
            info["db_error_detail"] = str(e)

    # Only expose paths to admin users
    if is_admin:
        current_db = _current_user_db_path()
        info["db_path"] = str(current_db)
        info["db_exists"] = current_db.exists()
        info["per_user_db"] = str(current_db) != str(DB_PATH)

    return jsonify(info)

@misc_bp.get("/api/data")
def api_data():
    uid = _uid()
    if uid is None:
        return jsonify(ok=False, error="Non authentifié"), 401
    # v23.4: Optional pagination via ?page=&limit= query params
    page_param = request.args.get("page")
    limit_param = request.args.get("limit")
    if page_param is not None:
        # Paginated mode
        try:
            page = max(1, int(page_param))
            limit = min(500, max(1, int(limit_param or 200)))
        except (TypeError, ValueError):
            return jsonify(ok=False, error="page/limit must be integers"), 400
        offset = (page - 1) * limit
        # v23.5: lazy=1 excludes heavy fields (callNotes, notes) for faster list loading
        lazy = request.args.get("lazy") == "1"
        with _conn() as conn:
            # Companies: always return all (typically small dataset)
            companies = [dict(r) for r in conn.execute(
                "SELECT * FROM companies WHERE owner_id=? AND deleted_at IS NULL ORDER BY id;", (uid,)
            ).fetchall()]
            # Prospects: paginated
            total = int(conn.execute(
                "SELECT COUNT(*) FROM prospects WHERE owner_id=? AND deleted_at IS NULL;", (uid,)
            ).fetchone()[0])
            prospects_rows = conn.execute(
                "SELECT * FROM prospects WHERE owner_id=? AND deleted_at IS NULL ORDER BY id LIMIT ? OFFSET ?;",
                (uid, limit, offset)
            ).fetchall()
            max_pid = int(conn.execute(
                "SELECT COALESCE(MAX(id), 0) FROM prospects WHERE owner_id=?;", (uid,)
            ).fetchone()[0])
            max_cid = int(conn.execute(
                "SELECT COALESCE(MAX(id), 0) FROM companies WHERE owner_id=?;", (uid,)
            ).fetchone()[0])
        # Parse tags/callNotes
        from math import ceil
        for c in companies:
            t = c.get("tags")
            if t and isinstance(t, str):
                try:
                    c["tags"] = json.loads(t)
                except Exception:
                    c["tags"] = [x.strip() for x in t.split(",") if x.strip()]
            elif not t:
                c["tags"] = []
        prospects = []
        for r in prospects_rows:
            d = dict(r)
            if lazy:
                # v23.5: exclude heavy fields for list view performance
                d.pop("callNotes", None)
                d.pop("notes", None)
            else:
                try:
                    d["callNotes"] = json.loads(d.get("callNotes") or "[]")
                except Exception:
                    d["callNotes"] = []
            t = d.get("tags")
            if t and isinstance(t, str):
                try:
                    d["tags"] = json.loads(t)
                except Exception:
                    d["tags"] = [x.strip() for x in t.split(",") if x.strip()]
            elif not t:
                d["tags"] = []
            d["is_archived"] = int(d.get("is_archived") or 0)
            prospects.append(d)
        return jsonify({
            "companies": companies,
            "prospects": prospects,
            "maxProspectId": max_pid,
            "maxCompanyId": max_cid,
            "pagination": {
                "page": page,
                "limit": limit,
                "total": total,
                "pages": ceil(total / limit) if limit else 1,
            }
        })
    # Non-paginated mode (backward compatible)
    payload = read_all(owner_id=uid)
    with _conn() as conn:
        payload["maxProspectId"] = int(conn.execute(
            "SELECT COALESCE(MAX(id), 0) AS n FROM prospects WHERE owner_id=?;",
            (uid,),
        ).fetchone()["n"])
        payload["maxCompanyId"] = int(conn.execute(
            "SELECT COALESCE(MAX(id), 0) AS n FROM companies WHERE owner_id=?;",
            (uid,),
        ).fetchone()["n"])
    return jsonify(payload)


@misc_bp.post("/api/save")
def api_save():
    chk = _require_same_origin()
    if chk:
        return chk
    data, err = validate_payload({})
    if err:
        return err
    try:
        upsert_all(data)
    except ValueError as e:
        return jsonify(ok=False, error=str(e)), 400
    except Exception as e:
        if app.config.get("TESTING"):
            err_msg = str(e) + "\n" + traceback.format_exc()
        else:
            err_msg = "Erreur lors de l'enregistrement."
        return jsonify(ok=False, error=err_msg), 500
    _auto_snapshot_if_needed()
    return jsonify({"ok": True})


def _excel_map_pertinence(val: str | None) -> str | None:
    if not val:
        return None
    s = str(val).strip()
    mapping = {
        "À contacter": "Pas d'actions",
        "A contacter": "Pas d'actions",
        "Appelé": "Appelé",
        "A rappeler": "À rappeler",
        "À rappeler": "À rappeler",
        "Rendez-vous": "Rendez-vous",
        "Prospecté": "Prospecté",
        "Messagerie": "Messagerie",
        "Pas intéressé": "Pas intéressé",
        "Pas interesse": "Pas intéressé",
    }
    return mapping.get(s, s)

def _excel_map_statut(val: str | None) -> str | None:
    if not val:
        return None
    s = str(val).strip()
    # Normaliser vers les libellés "simples" utilisés dans l'UI.
    mapping = {
        "À contacter": "Pas d'actions",
        "A contacter": "Pas d'actions",
        "□ Pas d'actions": "Pas d'actions",

        "Appelé": "Appelé",

        "A rappeler": "À rappeler",
        "À rappeler": "À rappeler",

        "Rendez-vous": "Rendez-vous",

        "Prospecté": "Prospecté",

        "Messagerie": "Messagerie",

        "Pas intéressé": "Pas intéressé",
        "Pas interesse": "Pas intéressé",
    }
    return mapping.get(s, s)

def _excel_cell_str(v):
    """Normalise les retours à la ligne pour Excel (évite \\r qui provoque décalages)."""
    if v is None or not isinstance(v, str):
        return v
    return v.replace("\r\n", "\n").replace("\r", "\n")


def _excel_concat_notes(prospect: dict) -> str | None:
    parts = []
    notes = (prospect.get("notes") or "").strip()
    if notes:
        parts.append(_excel_cell_str(notes))

    call_notes = prospect.get("callNotes") or []
    for n in call_notes:
        d = (n.get("date") or "").strip()
        c = (n.get("content") or "").strip()
        if not (d or c):
            continue
        if d and c:
            parts.append(f"{d} - {_excel_cell_str(c)}")
        else:
            parts.append(d or _excel_cell_str(c))
    out = "\n".join(parts).strip()
    return out or None


@misc_bp.get("/api/export/xlsx")
def api_export_xlsx():
    """Génère un fichier Excel à partir du template et des données SQLite (ligne entreprise + lignes prospects)."""
    from openpyxl import load_workbook

    if not TEMPLATE_PATH.exists():
        return jsonify({"ok": False, "error": "Template Excel introuvable"}), 500

    uid = _uid()
    if not uid:
        return jsonify(ok=False, error="Non authentifié"), 401
    payload = read_all(owner_id=uid)
    companies = payload.get("companies") or []
    prospects = payload.get("prospects") or []

    # Index prospects par entreprise
    pros_by_company: dict[int, list[dict]] = {}
    for p in prospects:
        try:
            cid = int(p.get("company_id"))
        except Exception:
            continue
        pros_by_company.setdefault(cid, []).append(p)

    # Tri : entreprises par groupe/site, prospects par nom
    companies_sorted = sorted(
        companies,
        key=lambda c: (str(c.get("groupe", "")).lower(), str(c.get("site", "")).lower()),
    )
    for cid in list(pros_by_company.keys()):
        pros_by_company[cid] = sorted(pros_by_company[cid], key=lambda p: str(p.get("name", "")).lower())

    wb = load_workbook(TEMPLATE_PATH)
    ws = wb["Liste"]

    # Headers sur la ligne 1
    headers: dict[str, int] = {}
    for c in range(1, ws.max_column + 1):
        v = ws.cell(1, c).value
        if isinstance(v, str) and v.strip():
            headers[v.strip()] = c

    # Styles sources (dans votre template) : une ligne entreprise + une ligne prospect
    company_style_row = 3
    prospect_style_row = 4
    max_col = ws.max_column

    def _capture_row_style(src_row: int):
        style = {"height": ws.row_dimensions[src_row].height, "cells": []}
        for col in range(1, max_col + 1):
            c = ws.cell(src_row, col)
            style["cells"].append(
                {
                    "_style": copy(c._style),
                    "font": copy(c.font),
                    "fill": copy(c.fill),
                    "border": copy(c.border),
                    "alignment": copy(c.alignment),
                    "number_format": c.number_format,
                    "protection": copy(c.protection),
                }
            )
        return style

    company_style = _capture_row_style(company_style_row)
    prospect_style = _capture_row_style(prospect_style_row)

    def _apply_row_style(style, dst_row: int):
        ws.row_dimensions[dst_row].height = style.get("height")
        for col in range(1, max_col + 1):
            d = ws.cell(dst_row, col)
            st = style["cells"][col - 1]
            d._style = copy(st["_style"])
            d.font = copy(st["font"])
            d.fill = copy(st["fill"])
            d.border = copy(st["border"])
            d.alignment = copy(st["alignment"])
            d.number_format = st["number_format"]
            d.protection = copy(st["protection"])
            d.comment = None

    def set_cell(row: int, header: str, value):
        col = headers.get(header)
        if not col:
            return
        if isinstance(value, str):
            value = _excel_cell_str(value)
        ws.cell(row, col).value = value

    def parse_date(iso: str | None):
        if not iso:
            return None
        try:
            return datetime.datetime.strptime(iso[:10], "%Y-%m-%d").date()
        except Exception:
            return None

    # Nettoyer anciennes données (garder lignes 1-2 du template)
    start_row = 3
    if ws.max_row >= start_row:
        ws.delete_rows(start_row, ws.max_row - start_row + 1)

    current_row = start_row

    for comp in companies_sorted:
        cid = int(comp["id"])

        # Ligne entreprise
        ws.insert_rows(current_row)
        _apply_row_style(company_style, current_row)
        set_cell(current_row, "GROUPE", comp.get("groupe"))
        set_cell(current_row, "SITE", comp.get("site"))
        set_cell(current_row, "TEL", comp.get("phone"))
        current_row += 1

        # Lignes prospects
        for p in pros_by_company.get(cid, []):
            ws.insert_rows(current_row)
            _apply_row_style(prospect_style, current_row)
            set_cell(current_row, "NOM", p.get("name"))
            set_cell(current_row, "TEL", p.get("telephone"))
            set_cell(current_row, "FONCTION", p.get("fonction"))
            set_cell(current_row, "PERTINENCE", _excel_map_pertinence(p.get("pertinence")))
            set_cell(current_row, "STATUT", _excel_map_statut(p.get("statut")))
            set_cell(current_row, "DATE DERNIER CONTACT", parse_date(p.get("lastContact")))
            set_cell(current_row, "COMMENTAIRE", _excel_concat_notes(p))
            set_cell(current_row, "MAIL", p.get("email"))
            set_cell(current_row, "LINKEDIN_URL", p.get("linkedin"))
            current_row += 1

    # Sauvegarde en mémoire et téléchargement
    bio = BytesIO()
    wb.save(bio)
    bio.seek(0)

    filename = f"Prospects_export_{datetime.date.today().isoformat()}.xlsx"
    return send_file(
        bio,
        as_attachment=True,
        download_name=filename,
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )

# ────────────────────────────────────────────────────────────────────
# Export "Ma journée" (P7) – récap du jour pour téléchargement
# ────────────────────────────────────────────────────────────────────

@misc_bp.get("/api/export/day")
def api_export_day():
    """Return a JSON recap of the day (contacts, notes, push, overdue, due_today) for download."""
    uid = _uid()
    if not uid:
        return jsonify(ok=False, error="Non authentifié"), 401
    date_str = request.args.get("date", "").strip() or _today_iso()
    try:
        datetime.date.fromisoformat(date_str)
    except Exception:
        date_str = _today_iso()

    with _conn() as conn:
        prospects = [dict(r) for r in conn.execute("SELECT * FROM prospects WHERE owner_id=?;", (uid,)).fetchall()]
        push_logs = [dict(r) for r in conn.execute(
            "SELECT * FROM push_logs WHERE prospect_id IN (SELECT id FROM prospects WHERE owner_id=?);",
            (uid,),
        ).fetchall()]
        try:
            note_events = [dict(r) for r in conn.execute(
                """SELECT e.date, e.content, e.prospect_id, p.name AS prospect_name
                   FROM prospect_events e
                   JOIN prospects p ON p.id=e.prospect_id
                   WHERE p.owner_id=? AND e.type IN ('note','note_libre','call_note');""",
                (uid,),
            ).fetchall()]
        except Exception:
            note_events = []

    all_notes = []
    for p in prospects:
        try:
            notes = json.loads(p.get("callNotes") or "[]")
            for n in (notes if isinstance(notes, list) else []):
                n["_pid"] = p["id"]
                n["_name"] = p.get("name", "")
                all_notes.append(n)
        except Exception:
            pass
    for ne in note_events:
        all_notes.append({
            "date": ne.get("date") or "",
            "content": ne.get("content") or "",
            "_pid": ne.get("prospect_id"),
            "_name": ne.get("prospect_name") or "",
        })

    contacts_today = [p for p in prospects if (p.get("lastContact") or "").strip() == date_str]
    notes_today = [n for n in all_notes if (n.get("date") or "")[:10] == date_str]
    push_today = [pl for pl in push_logs if (pl.get("sentAt") or "")[:10] == date_str]
    overdue = [p for p in prospects if (p.get("nextFollowUp") or "").strip() and p["nextFollowUp"].strip() < date_str]
    due_today = [p for p in prospects if (p.get("nextFollowUp") or "").strip() == date_str]

    recap = {
        "date": date_str,
        "relances_count": len(contacts_today),
        "relances": [{"id": p["id"], "name": p.get("name"), "company_id": p.get("company_id")} for p in contacts_today],
        "notes_count": len(notes_today),
        "notes": [{"prospect_id": n.get("_pid"), "prospect_name": n.get("_name"), "date": n.get("date"), "content": (n.get("content") or "")[:200]} for n in notes_today],
        "push_count": len(push_today),
        "push": [{"prospect_id": pl.get("prospect_id"), "subject": pl.get("subject"), "to_email": pl.get("to_email"), "sentAt": pl.get("sentAt")} for pl in push_today],
        "overdue_count": len(overdue),
        "due_today_count": len(due_today),
    }
    return jsonify(ok=True, recap=recap)


# ────────────────────────────────────────────────────────────────────
# Rapport hebdomadaire – export markdown / copie OneNote
# ────────────────────────────────────────────────────────────────────

@misc_bp.get("/rapport")
def page_rapport():
    return redirect("/v30/stats", code=302)


@misc_bp.get("/api/rapport-hebdo")
def api_rapport_hebdo():
    """Generate a weekly report with KPIs, activity, and pipeline summary."""
    # Determine week: defaults to current, or ?week=2026-W07
    week_param = request.args.get("week", "").strip()
    today = _today_iso()
    d_today = datetime.date.fromisoformat(today)

    if week_param:
        # Parse ISO week like 2026-W07
        try:
            year, w = week_param.split("-W")
            year, w = int(year), int(w)
            # Monday of that week
            jan4 = datetime.date(year, 1, 4)
            start_of_w1 = jan4 - datetime.timedelta(days=jan4.isoweekday() - 1)
            monday = start_of_w1 + datetime.timedelta(weeks=w - 1)
            sunday = monday + datetime.timedelta(days=6)
        except Exception:
            monday = d_today - datetime.timedelta(days=d_today.weekday())
            sunday = monday + datetime.timedelta(days=6)
    else:
        monday = d_today - datetime.timedelta(days=d_today.weekday())
        sunday = monday + datetime.timedelta(days=6)

    start = monday.isoformat()
    end = sunday.isoformat()
    week_label = f"S{monday.isocalendar()[1]} — {start} → {end}"

    uid = _uid()
    if not uid:
        return jsonify(ok=False, error="Non authentifié"), 401
    with _conn() as conn:
        prospects = [dict(r) for r in conn.execute("SELECT * FROM prospects WHERE owner_id=?;", (uid,)).fetchall()]
        push_logs = [dict(r) for r in conn.execute(
            "SELECT * FROM push_logs WHERE prospect_id IN (SELECT id FROM prospects WHERE owner_id=?);",
            (uid,),
        ).fetchall()]
        companies = [dict(r) for r in conn.execute("SELECT * FROM companies WHERE owner_id=?;", (uid,)).fetchall()]
        try:
            calls_row = conn.execute(
                "SELECT COUNT(*) AS n FROM call_logs WHERE owner_id=? AND date>=? AND date<=?;",
                (uid, start, end),
            ).fetchone()
            calls_count = int(calls_row["n"]) if calls_row else 0
        except Exception:
            calls_count = 0
        try:
            note_events = [dict(r) for r in conn.execute(
                """SELECT e.date, e.content, e.prospect_id, p.name AS prospect_name,
                          p.statut AS prospect_statut, p.company_id AS prospect_company_id
                   FROM prospect_events e
                   JOIN prospects p ON p.id=e.prospect_id
                   WHERE p.owner_id=? AND e.type IN ('note','note_libre','call_note');""",
                (uid,),
            ).fetchall()]
        except Exception:
            note_events = []

    # Parse call notes (callNotes JSON + prospect_events de type note)
    all_notes = []
    for p in prospects:
        try:
            notes = json.loads(p.get("callNotes") or "[]")
            for n in (notes if isinstance(notes, list) else []):
                n["_pid"] = p["id"]
                n["_pname"] = p.get("name", "")
                n["_statut"] = p.get("statut", "")
                n["_company_id"] = p.get("company_id")
                all_notes.append(n)
        except Exception:
            pass
    for ne in note_events:
        all_notes.append({
            "date": ne.get("date") or "",
            "content": ne.get("content") or "",
            "_pid": ne.get("prospect_id"),
            "_pname": ne.get("prospect_name") or "",
            "_statut": ne.get("prospect_statut") or "",
            "_company_id": ne.get("prospect_company_id"),
        })

    week_notes = [n for n in all_notes if start <= (n.get("date") or "")[:10] <= end]
    week_push = [pl for pl in push_logs if start <= (pl.get("sentAt") or "")[:10] <= end]
    week_relances = [p for p in prospects if start <= (p.get("lastContact") or "") <= end]

    push_email = sum(1 for pl in week_push if pl.get("channel") == "email")
    push_linkedin = sum(1 for pl in week_push if pl.get("channel") == "linkedin")

    # Status snapshot (BUG 17 : on regroupe les statuts vides sous "Autre" et on filtre les archivés/supprimés)
    statuts = {}
    for p in prospects:
        if p.get("deleted_at") or p.get("is_archived"):
            continue
        s = (p.get("statut") or "").strip()
        if not s or s.lower() == "inconnu":
            s = "Autre"
        statuts[s] = statuts.get(s, 0) + 1

    rdv_count = statuts.get("Rendez-vous", 0)
    total = len(prospects)

    # Overdue
    overdue = [p for p in prospects if (p.get("nextFollowUp") or "").strip() and p["nextFollowUp"].strip() < today]

    # New contacts this week (lastContact in range AND not before)
    new_relances_count = len(week_relances)

    # Companies touched this week + stats par entreprise
    prospects_by_id = {p["id"]: p for p in prospects}
    companies_map = {c["id"]: c for c in companies}

    week_company_ids = set()
    for n in week_notes:
        cid = n.get("_company_id")
        if cid:
            week_company_ids.add(cid)
    for pl in week_push:
        pid = pl.get("prospect_id")
        p = prospects_by_id.get(pid)
        if p:
            week_company_ids.add(p.get("company_id"))

    # Compter pushs et relances par company_id pour la semaine
    company_push_counts: dict = {}
    for pl in week_push:
        p = prospects_by_id.get(pl.get("prospect_id"))
        if p and p.get("company_id"):
            cid = p["company_id"]
            company_push_counts[cid] = company_push_counts.get(cid, 0) + 1

    company_relance_counts: dict = {}
    for p in week_relances:
        cid = p.get("company_id")
        if cid:
            company_relance_counts[cid] = company_relance_counts.get(cid, 0) + 1

    def _company_name(cid):
        c = companies_map.get(cid, {})
        return c.get("groupe") or c.get("site") or f"ID {cid}"

    top_companies = sorted(
        [
            {
                "name": _company_name(cid),
                "pushs": company_push_counts.get(cid, 0),
                "prospects": company_relance_counts.get(cid, 0),
            }
            for cid in week_company_ids if cid
        ],
        key=lambda x: -(x["pushs"] + x["prospects"]),
    )[:15]

    # Activity detail (BUG 16 : on inclut prospect_name partout)
    notes_detail = [{
        "prospect_id": n.get("_pid"),
        "prospect_name": n.get("_pname", ""),
        "name": n.get("_pname", ""),
        "statut": n.get("_statut", ""),
        "content": (n.get("content") or "")[:150],
        "date": n.get("date", ""),
    } for n in sorted(week_notes, key=lambda x: x.get("date", ""))]

    push_detail = [{
        "channel": pl.get("channel", ""),
        "date": (pl.get("sentAt") or "")[:10],
        "prospect_id": pl.get("prospect_id"),
        "prospect_name": (prospects_by_id.get(pl.get("prospect_id"), {}) or {}).get("name", ""),
    } for pl in sorted(week_push, key=lambda x: x.get("sentAt", ""))]

    # Conversion rate
    conversion_pct = round((rdv_count / total) * 100, 1) if total else 0

    return jsonify(ok=True, data={
        "week_label": week_label,
        "start": start,
        "end": end,
        "kpi": {
            "relances": new_relances_count,
            "notes": len(week_notes),
            "push_total": len(week_push),
            "push_email": push_email,
            "push_linkedin": push_linkedin,
            "rdv": rdv_count,
            "overdue": len(overdue),
            "conversion_pct": conversion_pct,
            "total_prospects": total,
            "companies_touched": len(top_companies),
            "calls": calls_count,
        },
        "statuts": statuts,
        "top_companies": top_companies,
        "touched_companies": [c["name"] for c in top_companies],
        "notes_detail": notes_detail[:20],
        "push_detail": push_detail[:20],
    })

