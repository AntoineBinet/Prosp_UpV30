#!/usr/bin/env python3
"""
Parse un CV (DOCX ou PDF) et retourne un dict structuré.
Ne lève jamais d'exception — dégradation gracieuse si champ absent.
"""
import os
import re

_DATE_RE = re.compile(
    r'\b(\d{2}/\d{4}|\d{4}[-/]\d{2,4}|'
    r'(?:jan|f[eé]v|mar|avr|mai|juin|juil|ao[uû]t?|sep|oct|nov|d[eé]c)'
    r'[\w.]*\s+\d{4})',
    re.IGNORECASE
)
_EXP_LINE_RE = re.compile(
    r'\(\s*\d+\s*(mois|an)|r[ôo]le\s*:|poste\s*:|mission\s*:|contexte\s*:',
    re.IGNORECASE
)
# Bullet chars à supprimer en début d'item
_BULLET_STRIP = re.compile(
    r'^[\u2022\u2023\u25e6\u2043\u2219\u25cf\u25cb\u2714\u2713'
    r'\u279e\u27a4\u25b6\u2715\u2716\u27a2\u2023\u203a'
    r'\*\->\u27a4\u27a6\u27a1\u2192\u21e8'
    r'➤✦❖●◆▶►•\s]+'
)


class CVParser:

    def parse(self, filepath: str) -> dict:
        ext = os.path.splitext(filepath)[1].lower()
        try:
            if ext == '.docx':
                return self._parse_docx(filepath)
            elif ext == '.pdf':
                return self._parse_pdf(filepath)
            else:
                return self._empty()
        except Exception as e:
            print(f'CVParser error: {e}')
            return self._empty()

    def _empty(self):
        return {
            'nom': '', 'prenom': '', 'titre_poste': '',
            'annees_experience': '',
            'email': '', 'telephone': '', 'localisation': '',
            'competences': [],
            'experiences': [],
            'formations': [],
            'langues': [],
            'certifications': []
        }

    # ── Extraction PDF ────────────────────────────────────────────────────────
    def _extract_pdf_text(self, filepath: str) -> str:
        try:
            import pypdfium2 as pdfium
            pdf = pdfium.PdfDocument(filepath)
            pages = []
            for i in range(len(pdf)):
                tp = pdf[i].get_textpage()
                pages.append(tp.get_text_range())
            return '\n'.join(pages)
        except Exception:
            pass
        try:
            from pypdf import PdfReader
            return '\n'.join(p.extract_text() or '' for p in PdfReader(filepath).pages)
        except Exception:
            pass
        return ''

    def _parse_docx(self, filepath: str) -> dict:
        from docx import Document
        doc = Document(filepath)
        texte = '\n'.join(p.text for p in doc.paragraphs if p.text.strip())
        return self._parse_text(texte)

    def _parse_pdf(self, filepath: str) -> dict:
        texte = self._extract_pdf_text(filepath)
        return self._parse_text(texte) if texte.strip() else self._empty()

    # ── Prétraitement des lignes ──────────────────────────────────────────────
    def _preprocess(self, raw_lignes: list) -> list:
        """
        1. Supprime les bullets seuls (•), numéros de page.
        2. Fusionne les labels fragmentés sur 2 lignes.
        """
        result = []
        i = 0
        while i < len(raw_lignes):
            l = raw_lignes[i]

            # Supprimer pagination et bullets seuls
            if re.match(r'^page\s+\d', l, re.I):
                i += 1; continue
            if re.match(r'^[•\-–—*]\s*$', l):
                i += 1; continue

            next_l = raw_lignes[i + 1] if i + 1 < len(raw_lignes) else ''

            # Fusionner "X de/du/des/d'" + suite (ex: "Domaines de" + "compétences")
            if next_l and re.search(r"\s+d[eu]s?$|\s+de$|\s+d'$", l) and len(l) < 22:
                result.append(l.rstrip() + ' ' + next_l.lstrip())
                i += 2; continue

            # Fusionner ligne se terminant par virgule + courte suite
            # (ex: "Langages, Outils," + "Normes")
            if next_l and l.endswith(',') and len(l) < 28 and len(next_l) < 22:
                result.append(l.rstrip() + ' ' + next_l.lstrip())
                i += 2; continue

            # Fusionner mot seul court + "d'..." (ex: "Secteurs" + "d'activités")
            # Normaliser les apostrophes courbes (U+2018/U+2019) pour la comparaison
            _next_norm = next_l.replace('\u2019', "'").replace('\u2018', "'")
            if (next_l and len(l) <= 12
                    and not re.search(r'[•\d–\-]', l)
                    and not l.startswith('•')
                    and (_next_norm.lower().startswith("d'") or _next_norm.lower().startswith("de ") or _next_norm.lower().startswith("du "))):
                result.append(l.rstrip() + ' ' + next_l.lstrip())
                i += 2; continue

            result.append(l)
            i += 1
        return result

    # ── Parse principal ───────────────────────────────────────────────────────
    def _parse_text(self, texte: str) -> dict:
        data = self._empty()
        texte = re.sub(r'[^\S\n]+', ' ', texte)
        raw = [l.strip() for l in texte.split('\n') if l.strip()]
        lignes = self._preprocess(raw)

        # Nom / prénom
        for l in lignes[:8]:
            if re.match(r'^page\s+\d', l, re.I):
                continue
            mots = l.split()
            if 1 <= len(mots) <= 5 and len(l) < 60 and not re.search(r'[@\d/\\–]', l):
                up_ratio = sum(1 for m in mots if m[0].isupper()) / len(mots)
                if up_ratio >= 0.5:
                    parties = l.rsplit(' ', 1)
                    if len(parties) == 2:
                        data['prenom'], data['nom'] = parties[0], parties[1]
                    else:
                        data['nom'] = l
                    break

        # Email
        m = re.search(r'[\w.\-+]+@[\w.\-]+\.\w{2,}', texte)
        if m: data['email'] = m.group()

        # Téléphone
        m = re.search(r'(\+?\d[\d\s.\-()\u00a0]{7,16}\d)', texte)
        if m: data['telephone'] = re.sub(r'\s+', ' ', m.group()).strip()

        # Années d'expérience
        m = re.search(r'(\d+)\s*an[s]?\s*d.exp[ée]rience', texte, re.IGNORECASE)
        if m: data['annees_experience'] = f"{m.group(1)} ans d'expérience"

        # Titre poste
        TITLE_PAT = (
            r'ing[ée]nieur|d[ée]veloppeur|developer|manager|chef|responsable|'
            r'technicien|consultant|analyste|architecte|directeur|coordinateur|'
            r'expert|lead|senior|junior|stagiaire|apprenti|scrum|product\s+owner'
        )
        for l in lignes[1:12]:
            if re.search(TITLE_PAT, l, re.I) and 5 < len(l) < 120 and not _DATE_RE.search(l):
                data['titre_poste'] = l
                break

        # Sections
        sections = self._split_sections(lignes)

        if 'competences' in sections:
            data['competences'] = self._parse_competences(sections['competences'])

        if 'experiences' in sections:
            data['experiences'] = self._parse_experiences(sections['experiences'])

        if 'formations' in sections:
            data['formations'] = self._parse_formations(sections['formations'])

        if 'langues' in sections:
            # Mots qui ne sont pas des noms de langues (éviter faux positifs depuis expériences)
            _NON_LANGUE = re.compile(
                r'^(stage|mission|secteur|projet|poste|logiciels?|outils?|contexte'
                r'|r[ée]alisation|responsable|client|dur[ée]e|p[ée]riode|lieu|ville'
                r'|entreprise|soci[ée]t[ée]|groupe|page)',
                re.I
            )
            for l in sections['langues']:
                if re.search(r'–|—|-', l):
                    parts = re.split(r'\s*[–—\-]\s*', l, 1)
                else:
                    parts = re.split(r'\s*:\s*', l, 1)
                _niveau = parts[1].strip() if len(parts) == 2 else ''
                if (len(parts) == 2
                        and 2 < len(parts[0]) < 25
                        and not _DATE_RE.search(parts[0])
                        and not _DATE_RE.search(_niveau)  # niveau ne contient pas de date
                        and not _NON_LANGUE.search(parts[0])
                        and len(parts[0].split()) <= 3   # max 3 mots (ex: "Anglais américain")
                        and len(_niveau) < 50):  # niveaux courts (B2, Courant, Natif…)
                    data['langues'].append({
                        'langue': parts[0].strip(),
                        'niveau': parts[1].strip()
                    })

        if 'certifications' in sections:
            data['certifications'] = [
                _BULLET_STRIP.sub('', l).strip()
                for l in sections['certifications']
                if len(l) > 3
            ]
            data['certifications'] = [c for c in data['certifications'] if c]

        return data

    # ── Découpage en sections ─────────────────────────────────────────────────
    def _split_sections(self, lignes: list) -> dict:
        # Titres ancrés : la ligne entière doit être un titre (pas un mot au
        # milieu d'une phrase comme « 18 ans d'expérience »).
        KEYWORDS = {
            'competences': (
                r'^(?:comp[ée]tences?(?:\s+(?:techniques?|fonctionnelles?|cl[ée]s?'
                r'|,?\s+outils?(?:\s+et\s+secteurs?)?))?'
                r'|skills?|savoir[\s\-]faire|expertise(?:\s+technique)?'
                r'|profil\s+technique)\s*:?$'
            ),
            'experiences': (
                r'^(?:exp[ée]riences?(?:\s+(?:professionnelles?|de\s+travail))?'
                r'|parcours(?:\s+professionnel)?'
                r'|historique(?:\s+professionnel)?'
                r'|missions?(?:\s+professionnelles?)?'
                r'|carri[eè]re|postes?\s+occup[ée]s?'
                r'|professional\s+experience|work\s+experience)\s*:?$'
            ),
            'formations': (
                r'^(?:formations?(?:\s+(?:et\s+)?(?:dipl[ôo]mes?|langues?|acad[ée]miques?))?'
                r'|dipl[ôo]mes?'
                r'|[ée]tudes?(?:\s+sup[ée]rieures?)?'
                r'|education|cursus)\s*:?$'
            ),
            'langues': (
                r'^(?:langues?(?:\s+(?:[ée]trang[eè]res?|parl[ée]es?|ma[îi]tris[ée]es?))?'
                r'|languages?)\s*:?$'
            ),
            'certifications': (
                r'^(?:certifications?|certificats?)\s*:?$'
            ),
        }
        # Exclure les lignes qui contiennent « ans d'expérience », qui ne sont
        # pas des titres de section.
        SKIP_INLINE = re.compile(
            r'\b(?:\d+\s*ans?|plusieurs)\s+d.exp[ée]rience|mission\s*:|r[ôo]le\s*:'
            r'|poste\s*:|contexte\s*:',
            re.I
        )

        sections: dict = {}
        current = None
        buf: list = []

        for l in lignes:
            # Retirer les préfixes bullets pour la détection de section
            l_test = _BULLET_STRIP.sub('', l).strip()
            # Un titre ne contient pas de date, pas de "ans d'expérience", etc.
            is_title_candidate = (
                len(l_test) <= 60
                and not _DATE_RE.search(l_test)
                and not SKIP_INLINE.search(l_test)
            )
            matched = None
            if is_title_candidate:
                for key, pat in KEYWORDS.items():
                    if re.match(pat, l_test, re.I):
                        matched = key
                        break
            if matched:
                if matched == current:
                    # Déjà dans cette section : label de sous-catégorie → ajouter au buf
                    buf.append(l)
                else:
                    if current and buf:
                        sections[current] = buf
                    current = matched
                    buf = []
                continue
            if current:
                buf.append(l)

        if current and buf:
            sections[current] = buf
        return sections

    # ── Parse section compétences ─────────────────────────────────────────────
    def _parse_competences(self, lignes: list) -> list:
        """
        Produit une liste de {categorie, items}.
        Chaque catégorie devient une rangée dans le tableau du dossier.

        Deux cas à gérer :
        1. Tableau : chaque ligne = "Label      Item1, Item2, Item3"
           (séparateur = 2+ espaces / tabulation)
        2. Liste : alternance de lignes d'en-tête (courtes, finissant par « : »
           ou correspondant à un label canonique Up) et de lignes d'items.
        """
        # Labels canoniques de tableau Up Technologies (après fusion)
        UP_LABEL_RE = re.compile(
            r'^(?:domaines?\s+de\s+comp[ée]tences?'
            r'|langages?(?:,?\s+outils?)?(?:,?\s+normes?)?'
            r'|secteurs?\s+d.activit[ée]s?'
            r'|outils?(?:\s+et\s+technologies?)?'
            r'|comp[ée]tences?\s+(?:techniques?|fonctionnelles?|cl[ée]s?)'
            r'|technologies?(?:\s+et\s+environnement)?'
            r'|environnement\s+technique'
            r'|m[ée]thodes?(?:\s+et\s+outils?)?'
            r'|bases?\s+de\s+donn[ée]es?|bdd|sgbd'
            r'|frameworks?'
            r'|syst[èe]mes?\s+(?:d.exploitation)?|os'
            r'|cloud|devops|web|conception|divers'
            r'|tech(?:nologies?)?\s+xml|java(?:/j2ee)?'
            r')\s*:?$',
            re.I
        )
        # Noms à ne JAMAIS considérer comme catégorie (viennent des expériences
        # mal découpées). Liste noire : outils, clients connus, etc.
        BAD_CAT_RE = re.compile(
            r'^(?:sql|db2|oracle|java|python|eclipse|maven|jira|docker|git|spring'
            r'|subversion|svn|jenkins|linux|unix|cobol|angular|react|node|vue'
            r'|hertz|natixis|bnp|bnp\s*paribas|edf|hsbc|mediapost|gmf|axa|sanofi|efs'
            r'|generali|cedicam|ina|societe\s*generale|sg|orange|bouygues|thales'
            r'|projet|mission|secteur|client|entreprise|soci[ée]t[ée]|groupe'
            r'|contexte|r[ôo]le|poste|fonction'
            r')\.?\s*$',
            re.I
        )

        competences = []
        cat = {'categorie': 'Compétences', 'items': []}

        def _flush():
            nonlocal cat
            if cat['items']:
                competences.append(cat)
            cat = {'categorie': 'Compétences', 'items': []}

        for raw in lignes:
            l = raw
            # Filtres parasites
            if re.match(r'^page\s+\d', l, re.I): continue
            if re.match(r'^[•\-–]\s*$', l): continue
            if re.match(r'^\(\d+\s*mois\)', l): continue
            if _DATE_RE.search(l) and len(l) < 60 and len(l.split()) <= 5: continue
            if _EXP_LINE_RE.search(l): continue
            if re.search(r'technologies?\s+et\s+outils?\s+utilis', l, re.I): continue
            if re.search(r'utilis[ée]s?\s+au\s+quotidien', l, re.I): continue

            stripped = _BULLET_STRIP.sub('', l).strip()
            if not stripped or len(stripped) <= 2:
                continue

            # Cas 1a : ligne tabulaire "LABEL   items..." (2+ espaces)
            m_tab = re.match(r'^([^\s:][^:]{2,45}?)[ \t]{2,}(.+)$', stripped)
            if m_tab:
                lbl = m_tab.group(1).strip().rstrip(':').strip()
                rest = m_tab.group(2).strip()
                if (UP_LABEL_RE.match(lbl) or lbl.endswith(':')
                        or (len(lbl.split()) <= 4 and len(lbl) <= 35
                            and not BAD_CAT_RE.match(lbl))):
                    _flush()
                    cat = {'categorie': lbl.rstrip(':').strip(), 'items': []}
                    for it in re.split(r'\s*[,;]\s*', rest):
                        it = it.strip()
                        if it and len(it) <= 200:
                            cat['items'].append(it)
                    continue

            # Cas 1b : ligne commence par un label canonique suivi d'items
            # (séparateur = simple espace, fréquent en PDF extraction)
            m_prefix = re.match(
                r'^(Java(?:/J2EE)?|Tech(?:nologies?)?\s+XML|Base\s+de\s+donn[ée]es'
                r'|Bases?\s+de\s+donn[ée]es?|BDD|SGBD'
                r'|Web|Conception|Divers|Cloud|DevOps|Frameworks?'
                r'|Langages?|Outils?|Secteurs?|M[ée]thodes?'
                r'|Certifications?|Environnement\s+technique'
                r'|Syst[èe]mes?\s+(?:d.exploitation)?|OS|Os)\b\s*[:\-–]?\s+(.+)$',
                stripped, re.I
            )
            if m_prefix and not UP_LABEL_RE.match(stripped):
                lbl = m_prefix.group(1).strip().rstrip(':').strip()
                rest = m_prefix.group(2).strip()
                if rest and len(rest) < 400 and not BAD_CAT_RE.match(lbl):
                    _flush()
                    cat = {'categorie': lbl, 'items': []}
                    for it in re.split(r'\s*[,;]\s*', rest):
                        it = it.strip(' .')
                        if it and len(it) <= 200:
                            cat['items'].append(it)
                    continue

            # Cas 2 : en-tête de catégorie
            is_header = False
            if UP_LABEL_RE.match(stripped):
                is_header = True
            elif stripped.endswith(':') or stripped.endswith(' :'):
                # Titre explicite « XXX : »
                if len(stripped) < 65 and not BAD_CAT_RE.match(stripped.rstrip(':').strip()):
                    is_header = True

            if is_header:
                label = stripped.rstrip(':').rstrip(' :').strip()
                if label and label.lower() != cat['categorie'].lower():
                    _flush()
                    cat = {'categorie': label, 'items': []}
                continue

            # Item de compétence
            item = stripped
            if _DATE_RE.search(item) and len(item.split()) <= 6: continue
            if len(item) > 200: continue
            cat['items'].append(item)

        _flush()

        # Nettoyage final : fusionner catégories "Compétences" vide avec la
        # suivante ; dédupliquer les items ; ignorer catégories dont le nom
        # ressemble à un item d'expérience.
        clean = []
        for c in competences:
            name = c.get('categorie', '').strip()
            items = c.get('items', [])
            if not items:
                continue
            if BAD_CAT_RE.match(name) and len(items) <= 1:
                continue
            # Dédup en conservant l'ordre
            seen = set()
            uniq = []
            for it in items:
                key = it.lower().strip()
                if key and key not in seen:
                    seen.add(key)
                    uniq.append(it)
            clean.append({'categorie': name, 'items': uniq[:15]})
        return clean

    # ── Parse section expériences ─────────────────────────────────────────────
    def _parse_experiences(self, lignes: list) -> list:
        """
        Heuristique pour CV tabulaires Up Technologies :
        - Une ligne d'entreprise (courte, majuscules / Title Case, pas de date)
          précède immédiatement une ligne contenant une date (« Depuis 09/2020 »,
          « 01/2012 au 07/2012 »).
        - Les lignes « Projet … : », « Rôle : », « Mission : », « Secteur : »
          sont des métadonnées du bloc en cours.
        - « (N mois) » seul → durée de l'expérience courante.
        - Les lignes « ❖ », « ➢ », « • », « o » sont des bullets.
        - Les lignes avec « Technologies et outils utilisés au quotidien »
          précèdent la liste des outils.
        """
        # Patterns
        DATE_HEAD = re.compile(
            r'^(?:depuis\s+)?\d{2}/\d{4}(?:\s*(?:au|[àa]|[-–—])\s*'
            r'(?:aujourd.hui|pr[ée]sent|(?:fin\s+)?(?:\w+\s+){0,2}\d{4}|\d{2}/\d{4}))?\s*$|'
            r'^\d{4}\s*[-–—à]\s*(?:\d{4}|pr[ée]sent|aujourd.hui)\s*$',
            re.I
        )
        DUREE_ONLY = re.compile(r'^\(\s*\d+\s*(?:mois|ans?|mo)\s*\)\s*$', re.I)
        PROJET_RE = re.compile(r'^projet(?:\s+(?:de|d.))?\s*[:\-–]\s*(.+)', re.I)
        SECTOR_RE = re.compile(r'^secteur\s*:\s*(.+)', re.I)
        ROLE_RE   = re.compile(r'^(?:r[ôo]le|poste|mission|fonction)\s*:\s*(.+)', re.I)
        DUREE_RE  = re.compile(r'^\(\s*(\d+)\s*(mois|ans?|mo)\s*\)', re.I)
        TOOLS_HEADER_RE = re.compile(
            r'technologies?\s+et\s+outils?\s+utilis|'
            r'utilis[ée]s?\s+au\s+quotidien|'
            r'^logiciels?\s*(?:/|et)?\s*outils?\s*:',
            re.I
        )
        # "Nom d'entreprise" : ligne courte, contient au moins une majuscule,
        # pas de date, pas un bullet, pas un mot-clé outil/rôle.
        def _is_company_line(s):
            s = s.strip()
            if not s or len(s) > 70:
                return False
            if _DATE_RE.search(s):
                return False
            if re.match(r'^[•➢❖◆►▸o\-–—*]', s):
                return False
            if not re.search(r'[A-Za-z]', s):
                return False
            if len(s.split()) > 8:
                return False
            # Commence par une majuscule (nom d'entité)
            if not re.match(r'^[A-Z0-9(]', s):
                return False
            # Pas un rôle / section / mot-clé
            if re.match(r'^(?:r[ôo]le|poste|mission|fonction|secteur|contexte|projet'
                        r'|client|dur[ée]e|environnement|technologies?|logiciels?|outils?'
                        r'|r[ée]alisations?|stack)\s*[:\-]', s, re.I):
                return False
            # Pas un mot-outil seul
            if re.match(r'^(?:SQL|DB2|Oracle|Java|Python|Eclipse|Maven|Jira|Docker'
                        r'|Git|Spring|SVN|Jenkins|Linux|Unix|COBOL)\.?\s*$', s, re.I):
                return False
            return True

        # Pré-traitement : fusionner les entreprises découpées sur 2 lignes.
        # Pattern typique : une ligne d'entreprise courte suivie d'une autre
        # ligne courte qui continue (parenthèse fermée, accent, etc.), puis
        # une ligne de date. On concatène 1 → 1+2 uniquement si 2 se ferme
        # par « ) » ou commence par une minuscule (continuation).
        merged = []
        i = 0
        L = lignes
        while i < len(L):
            cur = L[i].strip()
            nxt = L[i+1].strip() if i+1 < len(L) else ''
            nxt2 = L[i+2].strip() if i+2 < len(L) else ''
            # Entreprise wrap : "EFS (Etablissement" + "Français du Sang)"
            if (cur and nxt and nxt2
                    and '(' in cur and ')' in nxt
                    and not DATE_HEAD.match(cur) and not DATE_HEAD.match(nxt)
                    and (DATE_HEAD.match(nxt2) or DUREE_ONLY.match(nxt2))
                    and len(cur) < 40 and len(nxt) < 40):
                merged.append(cur + ' ' + nxt)
                i += 2; continue
            # Entreprise multi-mots : "BNP Paribas" / "INA (Institut National" + "Audiovisuel)"
            # Déjà géré par la parenthèse.
            merged.append(cur)
            i += 1
        L = merged

        # Fusion des dates coupées sur 2 lignes (ex. "03/2019 à fin " + "Septembre 2019")
        merged_d = []
        i = 0
        DATE_START = re.compile(r'^\d{2}/\d{4}\s*(?:au|[àa]|[-–—])\s+(?:fin\s+)?[a-zéèêA-Z]*\s*$')
        while i < len(L):
            cur = L[i]
            nxt = L[i+1] if i+1 < len(L) else ''
            if DATE_START.match(cur.strip()) and re.match(r'^[A-Za-zéèê]+\s+\d{4}\s*$', nxt.strip()):
                merged_d.append(cur.strip() + ' ' + nxt.strip())
                i += 2; continue
            merged_d.append(cur)
            i += 1
        L = merged_d

        # Fusion des « Rôle : » coupés sur 2 lignes (le second commence par
        # une majuscule courte, pas un bullet).
        merged2 = []
        i = 0
        while i < len(L):
            cur = L[i]
            nxt = L[i+1] if i+1 < len(L) else ''
            if (re.match(r'^(?:r[ôo]le|poste|fonction|mission)\s*:', cur, re.I)
                    and nxt and len(nxt) < 45
                    and re.match(r'^[A-Z]', nxt)
                    and not DATE_HEAD.match(nxt) and not DUREE_ONLY.match(nxt)
                    and not re.match(r'^(?:secteur|contexte|projet|environnement|technologies?'
                                     r'|logiciels?|outils?|r[ée]alisations?)\s*:', nxt, re.I)
                    and not re.match(r'^[•➢❖◆►▸o\-–—*]', nxt)):
                merged2.append(cur + ' ' + nxt)
                i += 2; continue
            merged2.append(cur)
            i += 1
        L = merged2

        # Détecter d'abord tous les indices de ligne qui démarrent une
        # nouvelle expérience (ligne de date DATE_HEAD).
        date_indices = [i for i, l in enumerate(L)
                        if DATE_HEAD.match(l.strip()) and len(l.strip()) < 80]

        experiences = []

        def _new_exp():
            return {
                'entreprise': '', 'dates': '', 'duree': '',
                'secteur': '', 'poste': '', 'titre_projet': '',
                'intro': '',
                'sous_missions': [{'titre': 'Réalisations', 'bullets': []}],
                'outils': ''
            }

        # Marqueurs de bullet explicites dans le texte PDF
        PRIMARY_BULLET = re.compile(r'^[❖•●◆▶►✦➤*]\s*')
        SECONDARY_BULLET = re.compile(r'^[➢▸▶►>→o]\s+')
        ENV_TECH_RE = re.compile(r'^environnement\s+technique\s*:?\s*$', re.I)
        END_PUNCT = re.compile(r'[.!?:…\)]\s*$')

        def _is_continuation(prev, cur):
            """cur continue prev si cur n'a pas de bullet marker,
            et prev ne finit pas par ponctuation forte."""
            if not prev:
                return False
            if PRIMARY_BULLET.match(cur) or SECONDARY_BULLET.match(cur):
                return False
            # Ligne courte commençant par minuscule = clairement une continuation
            if re.match(r'^[a-zà-ÿ(]', cur):
                return True
            # Prev ne finit pas par ponctuation forte → continuation
            if not END_PUNCT.search(prev):
                return True
            return False

        for k, didx in enumerate(date_indices):
            exp = _new_exp()
            exp['dates'] = L[didx].strip()

            # Remonter en arrière pour récupérer les lignes d'entreprise
            company_parts = []
            j = didx - 1
            start_prev = date_indices[k-1] if k > 0 else -1
            while j > start_prev and len(company_parts) < 3:
                cand = L[j].strip()
                if not cand:
                    j -= 1; continue
                if re.match(r'^page\s+\d', cand, re.I):
                    j -= 1; continue
                if _is_company_line(cand):
                    company_parts.insert(0, cand)
                    j -= 1
                    continue
                break
            if company_parts:
                exp['entreprise'] = ' '.join(company_parts).strip()

            # Fin du bloc
            end = date_indices[k+1] if k+1 < len(date_indices) else len(L)
            if k+1 < len(date_indices):
                next_didx = date_indices[k+1]
                exclude = 0
                jj = next_didx - 1
                while jj > didx and exclude < 3:
                    cand = L[jj].strip()
                    if not cand:
                        jj -= 1; continue
                    if _is_company_line(cand):
                        exclude += 1; jj -= 1; continue
                    break
                end = next_didx - exclude

            # ── Collecter les lignes du corps, en respectant les marqueurs ──
            body_lines = []
            for i in range(didx + 1, end):
                l = L[i].strip()
                if not l: continue
                if re.match(r'^page\s+\d', l, re.I): continue
                if re.match(r'^\d+\s*/\s*\d+\s*$', l): continue  # "3 / 7"
                body_lines.append(l)

            # ── Parser méta + intro + bullets + outils ──
            idx = 0
            while idx < len(body_lines):
                l = body_lines[idx]
                # Durée seule
                if DUREE_ONLY.match(l):
                    exp['duree'] = l
                    idx += 1; continue
                # Méta
                m = SECTOR_RE.match(l)
                if m:
                    exp['secteur'] = m.group(1).strip()
                    idx += 1; continue
                m = ROLE_RE.match(l)
                if m:
                    exp['poste'] = m.group(1).strip()
                    idx += 1; continue
                m = PROJET_RE.match(l)
                if m:
                    exp['titre_projet'] = m.group(1).strip()
                    idx += 1; continue
                break

            # ── Parser corps : intro + bullets + environnement technique (outils) ──
            current = None  # 'intro' | 'bullet' | 'tools'
            buf = []

            def _flush_buf():
                nonlocal buf, current
                if not buf:
                    return
                merged = ' '.join(buf).strip()
                merged = re.sub(r'\s+', ' ', merged)
                if current == 'intro':
                    if exp['intro']:
                        exp['intro'] += ' ' + merged
                    else:
                        exp['intro'] = merged
                elif current == 'bullet':
                    if merged:
                        exp['sous_missions'][0]['bullets'].append(merged)
                elif current == 'tools':
                    if merged:
                        exp['outils'] = (exp['outils'] + ', ' + merged).strip(', ')
                buf = []

            while idx < len(body_lines):
                l = body_lines[idx]
                idx += 1

                # "Environnement technique :" ou équivalent → bascule outils
                stripped_no_bullet = PRIMARY_BULLET.sub('', l).strip()
                if ENV_TECH_RE.match(stripped_no_bullet) or TOOLS_HEADER_RE.search(l):
                    _flush_buf()
                    current = 'tools'
                    # Si la même ligne contient ":" suivi de contenu
                    m_out = re.search(r':\s*(.+)$', l)
                    if m_out:
                        rest = m_out.group(1).strip()
                        if rest and not ENV_TECH_RE.match(stripped_no_bullet):
                            buf.append(rest); _flush_buf()
                    continue

                # Marqueur secondaire (➢ / o) en mode outils = item outil
                if current == 'tools' and SECONDARY_BULLET.match(l):
                    _flush_buf()
                    item = SECONDARY_BULLET.sub('', l).strip()
                    if item:
                        buf.append(item); _flush_buf()
                    continue

                # Marqueur primaire (❖ / •) = nouveau bullet
                if PRIMARY_BULLET.match(l):
                    _flush_buf()
                    current = 'bullet'
                    content = PRIMARY_BULLET.sub('', l).strip()
                    if content:
                        buf.append(content)
                    continue

                # Secondaire en mode bullets → sous-item (on colle au bullet actuel
                # préfixé de « – » pour ne pas perdre l'info)
                if SECONDARY_BULLET.match(l) and current == 'bullet':
                    content = SECONDARY_BULLET.sub('', l).strip()
                    if content:
                        # Ajouter comme nouveau bullet avec prefix
                        _flush_buf()
                        buf.append(content)
                        _flush_buf()
                    continue

                # Pas de marqueur : continuation du buf ou intro initiale
                if current is None:
                    current = 'intro'
                # Continuation : ajouter au buf
                if current in ('intro', 'bullet', 'tools'):
                    # Si la ligne est elle-même une métadonnée (Rôle tardif), skip
                    if ROLE_RE.match(l) or SECTOR_RE.match(l):
                        continue
                    buf.append(l)

            _flush_buf()

            if exp['entreprise'] or exp['intro'] or exp['sous_missions'][0]['bullets']:
                experiences.append(exp)

        # Filtrer les expériences vides (sans entreprise ni titre ni bullet)
        filtered = []
        for e in experiences:
            has_content = (
                e.get('entreprise') or e.get('titre_projet') or e.get('dates')
                or any(e.get('sous_missions', [{}])[0].get('bullets', []))
            )
            if has_content:
                # Fallback titre_projet : privilégier le poste seul
                # (l'entreprise apparaît déjà dans le sous-titre du DC).
                if not e['titre_projet']:
                    if e['poste']:
                        e['titre_projet'] = e['poste']
                    elif e['entreprise']:
                        e['titre_projet'] = e['entreprise']
                else:
                    # Si titre_projet contient "chez {entreprise}", retirer
                    if e['entreprise']:
                        e['titre_projet'] = re.sub(
                            r'\s+chez\s+' + re.escape(e['entreprise']) + r'.*$',
                            '', e['titre_projet'], flags=re.I
                        ).strip()
                filtered.append(e)
        return filtered

    # ── Parse section formations ──────────────────────────────────────────────
    def _parse_formations(self, lignes: list) -> list:
        formations = []
        for l in lignes:
            if re.match(r'^page\s+\d', l, re.I): continue
            if len(l) < 4: continue
            # Ignorer les sous-titres "Formation" et "Langues" seuls
            if re.match(r'^(formation|langues?|dipl[ôo]mes?)\s*$', l, re.I): continue
            annee_m = re.search(r'(20\d{2}|19\d{2})', l)
            label = 'Certification' if re.search(r'certif|brevet|aws|azure|pmp|prince2|itil|cisco', l, re.I) else 'Formation'
            formations.append({'label': label, 'texte': l, 'annee': annee_m.group(1) if annee_m else ''})
        return formations

    # ── Merge candidat DB ─────────────────────────────────────────────────────
    def merge_with_candidate(self, cv_data: dict, candidate_db: dict) -> dict:
        merged = dict(cv_data)
        for key in ['nom', 'prenom', 'titre_poste', 'email', 'telephone', 'localisation']:
            db_val = candidate_db.get(key, '')
            if db_val and str(db_val).strip():
                merged[key] = db_val
        if not merged.get('annees_experience'):
            for field in ('annees_experience', 'experience_years', 'years_experience'):
                yrs = candidate_db.get(field)
                if yrs:
                    merged['annees_experience'] = f"{yrs} ans d'expérience"
                    break
        return merged
