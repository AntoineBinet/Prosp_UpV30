#!/usr/bin/env python3
"""
DossierGenerator — Remplissage du template template_dc.docx.

Approche : ouvre le template tel quel, remplace les {{placeholders}},
duplique le bloc {{#EXPERIENCES}}…{{/EXPERIENCES}} pour chaque mission.
Mise en page Word préservée à 100%.

Placeholders du template :
  {{PRENOM_NOM}}, {{TITRE_POSTE}}, {{ANNEES_EXPERIENCE}}
  {{COMPETENCES}}, {{OUTILS}}, {{SECTEURS}}
  {{FORMATIONS}}, {{Année d'obtention}}, {{LANGUES}}
  {{#EXPERIENCES}} … {{/EXPERIENCES}} avec :
    {{EXP_TITRE}}, {{EXP_SECTEUR}}, {{EXP_MISSION}},
    {{EXP_REALISATIONS}}, {{EXP_OUTILS}}, {{Logo société}}
"""

import copy
import os
import re

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

_HERE = os.path.dirname(os.path.abspath(__file__))
TEMPLATE_PATH = os.path.join(_HERE, '..', 'sample', 'template_dc.docx')


# ── Utilitaires ───────────────────────────────────────────────────────────────

def _s(v) -> str:
    return str(v).strip() if v else ''


def _get_para_text(elem) -> str:
    """Texte complet d'un élément XML en joignant tous les w:t."""
    return ''.join(t.text or '' for t in elem.iter(qn('w:t')))


# ── Formatage des données ─────────────────────────────────────────────────────

def _fmt_prenom_nom(data: dict) -> str:
    prenom = _s(data.get('prenom'))
    nom    = _s(data.get('nom')).upper()
    return f"{prenom} {nom}".strip()


def _fmt_annees(data: dict) -> str:
    """Extrait juste le chiffre (le template ajoute ' ans d'expérience')."""
    val = _s(data.get('annees_experience'))
    m = re.match(r'^(\d+)', val)
    return m.group(1) if m else val


def _categorize_competences(competences: list) -> tuple:
    """
    Répartit les catégories en (domaines, outils, secteurs).
    Anti-duplication : si domaines vide, toutes les cats vont dans domaines
    et outils reste vide.
    """
    domaines, outils_cats, secteurs_cats = [], [], []

    OUTILS_KEYS  = ('outil', 'langag', 'framework', 'technolog', 'cloud',
                    'base de donn', 'web', 'normes', 'logiciel', 'infrastructure',
                    'certification', 'méthode')
    SECTEUR_KEYS = ('secteur', 'industrie', 'activit')
    # DOMAINE_KEYS volontairement strict : on ne matche que les catégories clairement
    # "domaine métier" (avec groupes) pour éviter la duplication avec OUTILS
    DOMAINE_KEYS = ('domaine', 'compét', 'expertise', 'savoir', 'architectur')

    for cat in competences:
        name = _s(cat.get('categorie')).lower()
        if any(k in name for k in SECTEUR_KEYS):
            secteurs_cats.append(cat)
        elif any(k in name for k in DOMAINE_KEYS) or cat.get('groupes'):
            domaines.append(cat)
        elif any(k in name for k in OUTILS_KEYS):
            outils_cats.append(cat)
        else:
            # Par défaut : catégorie plate sans groupe → outils
            outils_cats.append(cat)

    # Anti-duplication : si pas de domaines clairement identifiés,
    # tout va dans domaines (cellule COMPETENCES) et on vide outils
    if not domaines:
        domaines  = list(outils_cats)
        outils_cats = []

    return domaines, outils_cats, secteurs_cats


def _competences_to_lines(cats: list) -> list:
    """Retourne [(text, is_bold)] : titre de groupe en gras, items en bullet."""
    lines = []
    for cat in cats:
        name = _s(cat.get('categorie'))
        if cat.get('groupes'):
            lines.append((name, True))
            for g in cat['groupes']:
                titre = _s(g.get('titre'))
                items = [_s(x) for x in (g.get('items') or []) if x]
                if titre:
                    lines.append((titre, False))
                lines.extend((f'\u2022 {item}', False) for item in items)
        elif cat.get('items'):
            lines.append((name, True))
            lines.extend((f'\u2022 {_s(x)}', False) for x in cat['items'] if x)
    return lines


def _outils_to_lines(cats: list) -> list:
    """Retourne [(text, is_bold)] pour les outils (bullet par catégorie)."""
    lines = []
    for cat in cats:
        name  = _s(cat.get('categorie'))
        items = []
        for x in (cat.get('items') or []):
            if x: items.append(_s(x))
        for g in (cat.get('groupes') or []):
            for x in (g.get('items') or []):
                if x: items.append(_s(x))
        if items:
            lines.append((f'{name} : {", ".join(items)}', False))
    return lines


def _secteurs_to_lines(cats: list) -> list:
    lines = []
    items = []
    for cat in cats:
        for x in (cat.get('items') or []):
            if x: items.append(_s(x))
        for g in (cat.get('groupes') or []):
            for x in (g.get('items') or []):
                if x: items.append(_s(x))
    if items:
        lines.extend((f'\u2022 {item}', False) for item in items)
    return lines


def _formations_to_lines(formations: list) -> list:
    """Texte de formation sans le préfixe label (la colonne gauche l'indique déjà)."""
    lines = []
    for f in formations:
        texte = _s(f.get('texte'))
        if texte:
            lines.append((texte, False))
    return lines


def _annees_to_lines(formations: list) -> list:
    """Une année par formation (aligné avec _formations_to_lines)."""
    return [(_s(f.get('annee')), False) for f in formations]


def _langues_to_text(langues: list) -> str:
    parts = []
    for lang in langues:
        langue = _s(lang.get('langue'))
        niveau = _s(lang.get('niveau'))
        if langue:
            parts.append(f'{langue} ({niveau})' if niveau else langue)
    return ', '.join(parts)


# ── Formatage expériences ─────────────────────────────────────────────────────

def _fmt_exp_titre(exp: dict) -> str:
    titre      = _s(exp.get('titre_projet'))
    entreprise = _s(exp.get('entreprise'))
    dates      = _s(exp.get('dates'))
    duree      = _s(exp.get('duree'))
    parts = []
    if titre:      parts.append(titre)
    if entreprise: parts.append(f'\u2014 {entreprise}' if titre else entreprise)
    date_parts = [p for p in [dates, duree] if p]
    if date_parts: parts.append(f'({" ".join(date_parts)})')
    return ' '.join(parts)


def _fmt_exp_realisations(exp: dict) -> str:
    lines = []
    for sm in (exp.get('sous_missions') or []):
        bullets = list(sm.get('bullets') or [])
        for g in (sm.get('groupes') or []):
            g_titre = _s(g.get('titre'))
            g_items = [_s(x) for x in (g.get('items') or []) if x]
            if g_titre:
                lines.append(f'{g_titre} :')
                lines.extend(f'  \u2022 {x}' for x in g_items)
            else:
                bullets.extend(g_items)
        lines.extend(f'\u2022 {_s(b)}' for b in bullets if _s(b))
    return '\n'.join(lines)


# ── Manipulation XML ──────────────────────────────────────────────────────────

def _apply_bold(para_elem, bold: bool):
    """Active ou désactive le gras sur le premier run du paragraphe."""
    for r in para_elem.iter(qn('w:r')):
        rPr = r.find(qn('w:rPr'))
        if bold:
            if rPr is None:
                rPr = OxmlElement('w:rPr')
                r.insert(0, rPr)
            if rPr.find(qn('w:b')) is None:
                rPr.insert(0, OxmlElement('w:b'))
        else:
            if rPr is not None:
                b = rPr.find(qn('w:b'))
                if b is not None:
                    rPr.remove(b)
        break  # premier run seulement


def _set_para_text(para_elem, text: str):
    """Place le texte dans le premier w:t et vide les autres.
    Supporte le multiline via w:br."""
    t_elems = list(para_elem.iter(qn('w:t')))
    if not t_elems:
        return
    for t in t_elems[1:]:
        t.text = ''

    first_t = t_elems[0]

    if '\n' not in text:
        first_t.text = text
        if text and (text[0] == ' ' or text[-1] == ' '):
            first_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        return

    lines = text.split('\n')
    first_t.text = lines[0]
    first_r = first_t.getparent()
    if first_r is None:
        return
    for line in lines[1:]:
        first_r.append(OxmlElement('w:br'))
        t_new = OxmlElement('w:t')
        t_new.text = line
        if line and (line[0] == ' ' or line[-1] == ' '):
            t_new.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        first_r.append(t_new)


def _replace_in_para_elem(para_elem, replacements: dict):
    """
    Remplace les {{placeholders}} dans un w:p en gérant le fragment de runs.
    Supporte le texte multi-ligne (\\n → w:br).
    """
    full = _get_para_text(para_elem)
    if not any(k in full for k in replacements):
        return
    new_text = full
    for k, v in replacements.items():
        new_text = new_text.replace(k, _s(v))
    _set_para_text(para_elem, new_text)


def _expand_cell_content(doc: Document, placeholder: str, lines: list):
    """
    Trouve la cellule de tableau contenant le placeholder,
    remplace par plusieurs paragraphes (un par ligne).
    lines: [(text: str, bold: bool)]
    Les paragraphes sont des clones du paragraphe original (style préservé).
    """
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if placeholder not in _get_para_text(para._element):
                        continue
                    cell_tc  = cell._tc
                    para_elem = para._element
                    para_idx  = list(cell_tc).index(para_elem)

                    if not lines:
                        _replace_in_para_elem(para_elem, {placeholder: ''})
                        return

                    # Ligne 0 → modifier le paragraphe existant
                    text0, bold0 = lines[0]
                    _replace_in_para_elem(para_elem, {placeholder: text0})
                    _apply_bold(para_elem, bold0)

                    # Lignes suivantes → cloner + insérer
                    for j, (text, bold) in enumerate(lines[1:], 1):
                        new_p = copy.deepcopy(para_elem)
                        # Remettre text0 → text dans le clone
                        _set_para_text(new_p, text)
                        _apply_bold(new_p, bold)
                        cell_tc.insert(para_idx + j, new_p)
                    return


def _replace_in_all_paras(doc: Document, replacements: dict):
    """Remplace dans tous les w:p du document (body + tables + textboxes)."""
    for para in doc.element.body.iter(qn('w:p')):
        _replace_in_para_elem(para, replacements)


def _cleanup_placeholders(doc: Document):
    """
    Catch-all : supprime tout {{...}} résiduel dans l'ensemble du document.
    Gère à la fois les w:t atomiques et les runs fragmentés.
    """
    pattern = re.compile(r'\{\{[^}]*\}\}')
    for para in doc.element.body.iter(qn('w:p')):
        full = _get_para_text(para)
        if '{{' in full:
            new_text = pattern.sub('', full)
            _set_para_text(para, new_text)


def _expand_experiences(doc: Document, experiences: list):
    """
    Trouve {{#EXPERIENCES}}…{{/EXPERIENCES}} dans les enfants directs du body,
    supprime les marqueurs, et insère autant de copies du bloc qu'il y a d'expériences.
    Gère aussi les éléments non-paragraphe (tables, textboxes) dans le bloc.
    """
    body = doc.element.body
    start_elem = end_elem = None

    for child in body:
        if child.tag == qn('w:p') and '{{#EXPERIENCES}}' in _get_para_text(child):
            start_elem = child
        elif child.tag == qn('w:p') and '{{/EXPERIENCES}}' in _get_para_text(child) and start_elem is not None:
            end_elem = child
            break

    if start_elem is None or end_elem is None:
        return

    # Collecte des éléments du bloc (tous les enfants directs entre start et end)
    block_elems = []
    collecting  = False
    for child in list(body):
        if child is start_elem:
            collecting = True
            continue
        if child is end_elem:
            break
        if collecting:
            block_elems.append(child)

    insert_pos = list(body).index(start_elem)
    body.remove(start_elem)
    for elem in block_elems:
        body.remove(elem)
    body.remove(end_elem)

    for i, exp in enumerate(experiences):
        exp_replacements = {
            '{{EXP_TITRE}}':          _fmt_exp_titre(exp),
            '{{EXP_SECTEUR}}':        _s(exp.get('secteur')),
            '{{EXP_MISSION}}':        _s(exp.get('poste')),
            '{{EXP_REALISATIONS}}':   _fmt_exp_realisations(exp),
            '{{EXP_OUTILS}}':         _s(exp.get('outils')),
            '{{Logo soci\u00e9t\u00e9}}': '',   # placeholder logo → vide
        }
        offset = i * len(block_elems)
        for j, elem in enumerate(block_elems):
            new_elem = copy.deepcopy(elem)
            # Remplacer dans TOUS les w:p de l'élément (y compris tables internes)
            if new_elem.tag == qn('w:p'):
                _replace_in_para_elem(new_elem, exp_replacements)
            else:
                for para in new_elem.iter(qn('w:p')):
                    _replace_in_para_elem(para, exp_replacements)
            body.insert(insert_pos + offset + j, new_elem)


# ── Générateur principal ──────────────────────────────────────────────────────

class DossierGenerator:
    def __init__(self, template_path: str = None):
        self.template_path = template_path or TEMPLATE_PATH

    def generate(self, data: dict, output_path: str) -> str:
        """
        Remplit le template avec *data* et sauvegarde le .docx en *output_path*.
        *data* suit la structure de utils/ollama_extractor.py.
        """
        if not os.path.exists(self.template_path):
            raise FileNotFoundError(f'Template introuvable : {self.template_path}')

        doc = Document(self.template_path)

        competences = data.get('competences') or []
        experiences = data.get('experiences') or []
        formations  = data.get('formations')  or []
        langues     = data.get('langues')     or []

        domaines, outils_cats, secteurs_cats = _categorize_competences(competences)

        # Phase 1 : Bloc expériences (avant tout autre remplacement)
        _expand_experiences(doc, experiences)

        # Phase 2 : Cellules des tableaux (multi-paragraphes)
        _expand_cell_content(doc, '{{COMPETENCES}}',
                             _competences_to_lines(domaines))
        _expand_cell_content(doc, '{{OUTILS}}',
                             _outils_to_lines(outils_cats))
        _expand_cell_content(doc, '{{SECTEURS}}',
                             _secteurs_to_lines(secteurs_cats))
        _expand_cell_content(doc, '{{FORMATIONS}}',
                             _formations_to_lines(formations))
        _expand_cell_content(doc, '{{Ann\u00e9e d\u2019obtention}}',
                             _annees_to_lines(formations))
        langues_text = _langues_to_text(langues)
        _expand_cell_content(doc, '{{LANGUES}}',
                             [(langues_text, False)] if langues_text else [])

        # Phase 3 : Remplacements simples (nom, titre, années)
        _replace_in_all_paras(doc, {
            '{{PRENOM_NOM}}':        _fmt_prenom_nom(data),
            '{{TITRE_POSTE}}':       _s(data.get('titre_poste')),
            '{{ANNEES_EXPERIENCE}}': _fmt_annees(data),
        })

        # Phase 4 : Nettoyage de tous les {{...}} résiduels
        _cleanup_placeholders(doc)

        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        doc.save(output_path)
        return output_path
